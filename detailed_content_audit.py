from docx import Document
import re

book_path = r'public\reference\book\comptia-book.docx'

print("="*80)
print("DETAILED CONTENT AUDIT - PRACTICE QUESTIONS, PBQs, GLOSSARY")
print("="*80)

doc = Document(book_path)

# ============================================================================
print("\n" + "="*80)
print("PRACTICE TEST QUESTIONS - DETAILED VALIDATION")
print("="*80)

practice_start = -1
for i, para in enumerate(doc.paragraphs):
    if '📝 Practice Test' in para.text:
        practice_start = i
        break

if practice_start < 0:
    print("❌ Practice Test section not found")
else:
    print(f"✓ Practice Test starts at paragraph {practice_start}")
    
    # Extract all questions
    questions = []
    answer_key_start = -1
    
    for i in range(practice_start, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        
        # Look for answer key
        if ('Answer Key' in text or '🔑' in text) and 'Answer' in text:
            answer_key_start = i
            print(f"✓ Answer Key starts at paragraph {i}")
            break
        
        # Detect questions (usually followed by multiple choice A, B, C, D)
        if text and not text.startswith('A)') and not text.startswith('B)'):
            # Check if next few paragraphs have answer choices
            has_choices = False
            for j in range(i+1, min(i+6, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                if next_text.startswith('A)') or next_text.startswith('B)') or next_text.startswith('C)') or next_text.startswith('D)'):
                    has_choices = True
                    break
            
            if has_choices and '?' in text:
                questions.append({
                    'paragraph': i,
                    'question': text[:150]
                })
    
    print(f"\n✓ Found {len(questions)} practice questions")
    
    # Validate answer key
    if answer_key_start > 0:
        answers = []
        explanations = []
        
        for i in range(answer_key_start, min(answer_key_start + 500, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            
            # Look for answers (typically numbered or with answer pattern)
            if text and ('Answer' in text or 'Explanation' in text or 
                        text.startswith('A)') or text.startswith('B)') or 
                        text.startswith('C)') or text.startswith('D)')):
                answers.append(i)
        
        print(f"✓ Found {len(answers)} answer entries in answer key section")
        
        if len(questions) > len(answers):
            print(f"⚠️  WARNING: More questions ({len(questions)}) than answers ({len(answers)})")
        elif len(answers) > len(questions) * 2:
            print(f"⚠️  WARNING: Many more answer entries than questions - may include explanations (OK)")
        else:
            print(f"✓ Question/Answer count appears balanced")
    
    # Show sample questions
    print(f"\n📝 Sample Questions (first 5):")
    for i, q in enumerate(questions[:5], 1):
        print(f"\n  Q{i} (Paragraph {q['paragraph']}):")
        print(f"    {q['question']}")

# ============================================================================
print("\n" + "="*80)
print("PBQ/SKILLS TEST SCENARIOS - DETAILED VALIDATION")
print("="*80)

pbq_start = -1
for i, para in enumerate(doc.paragraphs):
    if ('Skills Test' in para.text and '🧪' in para.text) or 'Performance-Based Scenarios' in para.text:
        pbq_start = i
        break

if pbq_start < 0:
    print("❌ PBQ section not found")
else:
    print(f"✓ PBQ section starts at paragraph {pbq_start}")
    
    # Find all scenarios
    scenarios = []
    current_scenario = None
    
    for i in range(pbq_start, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        
        # Detect scenario headers
        if 'Scenario' in text and ':' in text:
            if current_scenario:
                scenarios.append(current_scenario)
            
            current_scenario = {
                'title': text,
                'paragraph': i,
                'has_situation': False,
                'has_task': False,
                'has_answer_key': False,
                'questions': []
            }
        
        elif current_scenario:
            if 'SITUATION' in text.upper():
                current_scenario['has_situation'] = True
            if 'YOUR TASK' in text.upper() or 'TASK' in text.upper():
                current_scenario['has_task'] = True
            if 'ANSWER KEY' in text.upper() and '✅' in text:
                current_scenario['has_answer_key'] = True
            if text.startswith('Q') and any(char in text for char in ['?', ':']):
                current_scenario['questions'].append(text[:80])
    
    if current_scenario:
        scenarios.append(current_scenario)
    
    print(f"\n✓ Found {len(scenarios)} PBQ scenarios")
    
    # Validate each scenario
    print(f"\n📊 PBQ Scenario Validation:")
    for i, scenario in enumerate(scenarios, 1):
        print(f"\n  Scenario {i}: {scenario['title'][:60]}")
        print(f"    Location: Paragraph {scenario['paragraph']}")
        print(f"    Has SITUATION: {'✓' if scenario['has_situation'] else '❌'}")
        print(f"    Has TASK: {'✓' if scenario['has_task'] else '❌'}")
        print(f"    Has ANSWER KEY: {'✓' if scenario['has_answer_key'] else '❌'}")
        print(f"    Questions: {len(scenario['questions'])}")
        
        if not scenario['has_situation']:
            print(f"    ⚠️  WARNING: Missing SITUATION section")
        if not scenario['has_task']:
            print(f"    ⚠️  WARNING: Missing YOUR TASK section")
        if not scenario['has_answer_key']:
            print(f"    ⚠️  WARNING: Missing ANSWER KEY")

# ============================================================================
print("\n" + "="*80)
print("GLOSSARY - DETAILED VALIDATION")
print("="*80)

glossary_start = -1
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if ('Glossary' in text or '📖' in text) and len(text) < 50:
        glossary_start = i
        break

if glossary_start < 0:
    print("❌ Glossary section not found")
else:
    print(f"✓ Glossary starts at paragraph {glossary_start}")
    
    # Extract glossary terms
    terms = []
    definitions = []
    current_term = None
    
    for i in range(glossary_start + 1, min(glossary_start + 5000, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        
        # Stop at next major section
        if any(marker in text for marker in ['📝', '🧪', '📊', '🧠']) and len(text) < 100:
            break
        
        if not text:
            continue
        
        # Check if this is a term (bold text or contains definition marker)
        para = doc.paragraphs[i]
        has_bold = any(run.bold for run in para.runs if run.text.strip())
        has_separator = ' — ' in text or ' – ' in text or ': ' in text
        
        if has_bold or has_separator:
            # Extract term
            if has_separator:
                parts = re.split(r'\s+[—–:]\s+', text, maxsplit=1)
                if len(parts) == 2:
                    term, definition = parts
                    terms.append({
                        'term': term.strip(),
                        'definition': definition.strip(),
                        'paragraph': i,
                        'length': len(definition.strip())
                    })
            elif has_bold:
                # Bold term, definition might be in same or next paragraph
                term_text = ''.join([run.text for run in para.runs if run.bold]).strip()
                if term_text and len(term_text) < 100:
                    definition_text = text.replace(term_text, '').strip()
                    terms.append({
                        'term': term_text,
                        'definition': definition_text if definition_text else '(see next paragraph)',
                        'paragraph': i,
                        'length': len(definition_text)
                    })
    
    print(f"\n✓ Found {len(terms)} glossary entries")
    
    # Validate glossary entries
    short_definitions = [t for t in terms if t['length'] < 20]
    long_terms = [t for t in terms if len(t['term']) > 100]
    
    if short_definitions:
        print(f"\n⚠️  Found {len(short_definitions)} entries with very short definitions (< 20 chars)")
        for t in short_definitions[:5]:
            print(f"    Paragraph {t['paragraph']}: {t['term'][:40]} — {t['definition'][:40]}")
    
    if long_terms:
        print(f"\n⚠️  Found {len(long_terms)} entries with very long terms (> 100 chars) - may be formatting issue")
        for t in long_terms[:5]:
            print(f"    Paragraph {t['paragraph']}: {t['term'][:60]}...")
    
    # Show sample entries
    print(f"\n📖 Sample Glossary Entries (first 10):")
    for t in terms[:10]:
        term_display = t['term'][:40]
        def_display = t['definition'][:60] if len(t['definition']) > 60 else t['definition']
        print(f"  {term_display} — {def_display}{'...' if len(t['definition']) > 60 else ''}")
    
    # Check for expected terms
    expected_terms = [
        'APIPA', 'DHCP', 'DNS', 'RAID', 'BIOS', 'UEFI', 'SSD', 'HDD',
        'IPv6', 'SLAAC', 'USB-C', 'Thread', 'Wi-Fi', 'VPN', 'Malware',
        'BitLocker', 'GPT', 'MBR', 'NTFS', 'FAT32'
    ]
    
    found_terms = [term for term in expected_terms if any(term.lower() in t['term'].lower() for t in terms)]
    missing_terms = [term for term in expected_terms if term not in found_terms]
    
    print(f"\n🔍 Expected Terms Check:")
    print(f"  Found: {len(found_terms)}/{len(expected_terms)}")
    if missing_terms:
        print(f"  Missing: {', '.join(missing_terms)}")
    else:
        print(f"  ✓ All expected terms present")

# ============================================================================
print("\n" + "="*80)
print("CHECKING FOR LEFTOVER INSERTION MARKERS")
print("="*80)

insertion_markers = [
    'INSERT HERE', 'TODO', 'PLACEHOLDER', 'TBD', 'XXX',
    '[insert', '[add ', 'FIXME', 'TEMP', '???'
]

found_markers = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.upper()
    for marker in insertion_markers:
        if marker in text:
            found_markers.append({
                'paragraph': i,
                'marker': marker,
                'text': doc.paragraphs[i].text[:100]
            })

if found_markers:
    print(f"⚠️  Found {len(found_markers)} potential insertion markers/placeholders:")
    for item in found_markers:
        print(f"  Paragraph {item['paragraph']}: {item['marker']}")
        print(f"    {item['text']}")
else:
    print("✓ No insertion markers or placeholders found")

# ============================================================================
print("\n" + "="*80)
print("SUMMARY")
print("="*80)

print(f"\n📊 Content Summary:")
print(f"  Total Paragraphs: {len(doc.paragraphs)}")
print(f"  Practice Questions: {len(questions) if practice_start >= 0 else 'Not found'}")
print(f"  PBQ Scenarios: {len(scenarios) if pbq_start >= 0 else 'Not found'}")
print(f"  Glossary Terms: {len(terms) if glossary_start >= 0 else 'Not found'}")

print(f"\n✅ CONTENT VALIDATION:")
if practice_start >= 0 and len(questions) > 20:
    print(f"  ✓ Practice Test: Present with {len(questions)} questions")
else:
    print(f"  ❌ Practice Test: Missing or incomplete")

if pbq_start >= 0 and len(scenarios) >= 5:
    print(f"  ✓ PBQ Scenarios: Present with {len(scenarios)} scenarios")
else:
    print(f"  ❌ PBQ Scenarios: Missing or incomplete")

if glossary_start >= 0 and len(terms) > 100:
    print(f"  ✓ Glossary: Present with {len(terms)} terms")
elif glossary_start >= 0:
    print(f"  ⚠️  Glossary: Present but appears incomplete ({len(terms)} terms)")
else:
    print(f"  ❌ Glossary: Not found")

print("\n" + "="*80)
print("DETAILED AUDIT COMPLETE")
print("="*80)
