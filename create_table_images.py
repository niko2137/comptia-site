"""
Create KDP-appropriate images for wide/complex tables
Generates clean, readable PNG images optimized for e-readers
"""

from docx import Document
from PIL import Image, ImageDraw, ImageFont
import os

SOURCE_FILE = r'public\reference\book\comptia-book.docx'
OUTPUT_DIR = r'public\reference\book\images'

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Tables to convert to images
TABLES_TO_IMAGE = {
    47: "Shared_Responsibility_Model.png",
    69: "Memory_Types_Comparison.png",
    73: "Display_Resolutions.png",
    82: "RAID_Configurations.png",
    91: "Windows_Versions.png",
    93: "Security_Protocols.png",
    133: "Cable_Specifications.png"
}

# Color scheme (Kindle-friendly: high contrast, grayscale-safe)
COLORS = {
    'bg': (255, 255, 255),           # White background
    'header_bg': (45, 45, 45),       # Dark gray header
    'header_text': (255, 255, 255),  # White text
    'border': (100, 100, 100),       # Medium gray borders
    'text': (20, 20, 20),            # Near-black text
    'alt_row': (245, 245, 245)       # Light gray alternating rows
}

def deduplicate_merged_cells(row):
    """Handle merged cells by deduplicating"""
    seen = set()
    unique = []
    for cell in row.cells:
        cell_id = id(cell._element)
        if cell_id not in seen:
            seen.add(cell_id)
            unique.append(cell)
    return unique

def extract_table_data(table):
    """Extract all data from table handling merged cells"""
    data = []
    for row in table.rows:
        cells = deduplicate_merged_cells(row)
        row_data = [cell.text.strip() for cell in cells]
        data.append(row_data)
    return data

def calculate_column_widths(data, min_width=80, max_width=400):
    """Calculate optimal column widths based on content"""
    if not data:
        return []
    
    num_cols = len(data[0])
    widths = [min_width] * num_cols
    
    for row in data:
        for i, cell in enumerate(row):
            # Estimate width based on text length
            text_length = len(cell)
            estimated_width = min(max_width, max(min_width, text_length * 8))
            widths[i] = max(widths[i], estimated_width)
    
    return widths

def wrap_text(text, font, max_width):
    """Wrap text to fit within max_width"""
    words = text.split()
    lines = []
    current_line = []
    
    for word in words:
        test_line = ' '.join(current_line + [word])
        bbox = font.getbbox(test_line)
        width = bbox[2] - bbox[0]
        
        if width <= max_width - 20:  # 20px padding
            current_line.append(word)
        else:
            if current_line:
                lines.append(' '.join(current_line))
            current_line = [word]
    
    if current_line:
        lines.append(' '.join(current_line))
    
    return lines

def create_table_image(table_data, filename, title=""):
    """Create a professional table image"""
    
    # Try to load fonts (fall back to default if not available)
    try:
        font_header = ImageFont.truetype("arialbd.ttf", 14)
        font_body = ImageFont.truetype("arial.ttf", 12)
    except:
        try:
            font_header = ImageFont.truetype("Arial.ttf", 14)
            font_body = ImageFont.truetype("Arial.ttf", 12)
        except:
            font_header = ImageFont.load_default()
            font_body = ImageFont.load_default()
    
    # Calculate dimensions
    col_widths = calculate_column_widths(table_data)
    row_height = 40
    header_height = 50
    padding = 10
    
    # Calculate image dimensions
    img_width = sum(col_widths) + (len(col_widths) + 1) * padding
    
    # Calculate row heights (accounting for text wrapping)
    row_heights = [header_height]
    for row in table_data[1:]:
        max_height = row_height
        for i, cell in enumerate(row):
            wrapped = wrap_text(cell, font_body, col_widths[i])
            cell_height = len(wrapped) * 18 + 20
            max_height = max(max_height, cell_height)
        row_heights.append(max_height)
    
    img_height = sum(row_heights) + padding * 2
    
    # Create image
    img = Image.new('RGB', (img_width, img_height), COLORS['bg'])
    draw = ImageDraw.Draw(img)
    
    # Draw table
    y = padding
    
    for row_idx, row in enumerate(table_data):
        x = padding
        is_header = (row_idx == 0)
        row_h = row_heights[row_idx]
        
        # Alternating row background
        if not is_header and row_idx % 2 == 0:
            draw.rectangle([padding, y, img_width - padding, y + row_h], 
                          fill=COLORS['alt_row'])
        
        for col_idx, cell in enumerate(row):
            col_w = col_widths[col_idx]
            
            # Cell background (header)
            if is_header:
                draw.rectangle([x, y, x + col_w, y + row_h], 
                             fill=COLORS['header_bg'])
            
            # Cell border
            draw.rectangle([x, y, x + col_w, y + row_h], 
                          outline=COLORS['border'], width=1)
            
            # Cell text
            text_color = COLORS['header_text'] if is_header else COLORS['text']
            font = font_header if is_header else font_body
            
            # Wrap text
            wrapped_lines = wrap_text(cell, font, col_w)
            
            # Draw wrapped text
            text_y = y + 10
            for line in wrapped_lines:
                draw.text((x + 10, text_y), line, fill=text_color, font=font)
                text_y += 18
            
            x += col_w
        
        y += row_h
    
    # Save image
    output_path = os.path.join(OUTPUT_DIR, filename)
    img.save(output_path, 'PNG', dpi=(300, 300), optimize=True)
    print(f"   ✓ Created: {filename} ({img_width}x{img_height}px)")
    
    return output_path

def main():
    print("="*80)
    print("CREATING TABLE IMAGES FOR KDP")
    print("="*80)
    
    # Load document
    print(f"\n📖 Loading document...")
    doc = Document(SOURCE_FILE)
    print(f"   ✓ {len(doc.tables)} tables found")
    
    # Create images
    print(f"\n🎨 Creating {len(TABLES_TO_IMAGE)} table images...")
    
    created_images = []
    
    for table_idx, filename in TABLES_TO_IMAGE.items():
        if table_idx >= len(doc.tables):
            print(f"   ⚠️  Table {table_idx} not found (only {len(doc.tables)} tables)")
            continue
        
        table = doc.tables[table_idx]
        data = extract_table_data(table)
        
        # Get table title from first cell
        title = data[0][0] if data else ""
        
        print(f"\n   Table {table_idx}: {len(data)}x{len(data[0]) if data else 0}")
        print(f"   Title: {title[:60]}")
        
        try:
            image_path = create_table_image(data, filename, title)
            created_images.append(image_path)
        except Exception as e:
            print(f"   ❌ Error creating image: {e}")
    
    # Summary
    print("\n" + "="*80)
    print("✅ IMAGE CREATION COMPLETE")
    print("="*80)
    print(f"Created {len(created_images)} images:")
    for img in created_images:
        print(f"   ✓ {os.path.basename(img)}")
    
    print(f"\n📁 Images saved to: {OUTPUT_DIR}")
    print("\n✅ Ready for KDP conversion script!")

if __name__ == '__main__':
    main()
