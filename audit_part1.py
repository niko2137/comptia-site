"""Part 1: Structure, formatting, spacing audit - KDP only first"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from collections import defaultdict

print("Loading KDP-final.docx...")
doc = Document(r'public\reference\book\KDP-final.docx')
print(f"Loaded: {len(doc.paragraphs)} paragraphs")

out = open('audit_p1_results.txt', 'w', encoding='utf-8')

def w(text=""):
    out.write(text + "\n")

w("PART 1: STRUCTURE & FORMATTING - KDP-final.docx")
w("=" * 60)
w(f"Paragraphs: {len(doc.paragraphs)}")
w(f"Tables: {len(doc.tables)}")
w(f"Sections: {len(doc.sections)}")

# 1. Tables
w(f"\n[Tables] {'PASS' if not doc.tables else 'FAIL'}: {len(doc.tables)} tables")

# 2. Page breaks
pb = 0
for para in doc.paragraphs:
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:pageBreakBefore')) is not None:
        pb += 1
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                pb += 1
w(f"[Page Breaks] {'PASS' if pb == 0 else 'FAIL'}: {pb}")

# 3. Headers/Footers
hf = 0
for s in doc.sections:
    for p in s.header.paragraphs:
        if p.text.strip(): hf += 1
    for p in s.footer.paragraphs:
        if p.text.strip(): hf += 1
w(f"[Headers/Footers] {'PASS' if hf == 0 else 'FAIL'}: {hf} with text")

# 4. Double spaces
ds = [(i, p.text[:60]) for i, p in enumerate(doc.paragraphs) if '  ' in p.text]
w(f"[Double Spaces] {'PASS' if not ds else 'WARN'}: {len(ds)} paragraphs")
for idx, txt in ds[:5]:
    w(f"  Para {idx}: {txt}")

# 5. Leading whitespace
lw = [(i, p.text[:50]) for i, p in enumerate(doc.paragraphs) if p.text and p.text[0] in ' \t']
w(f"[Leading Whitespace] {'PASS' if not lw else 'WARN'}: {len(lw)} paragraphs")
for idx, txt in lw[:5]:
    w(f"  Para {idx}: '{txt}'")

# 6. Consecutive blanks
max_b = 0
cur = 0
for p in doc.paragraphs:
    has_img = bool(p._element.findall('.//' + qn('w:drawing')))
    if not p.text.strip() and not has_img:
        cur += 1
        max_b = max(max_b, cur)
    else:
        cur = 0
w(f"[Blank Runs] Max consecutive: {max_b}")

# 7. Hidden chars
hid = defaultdict(int)
for p in doc.paragraphs:
    for r in p.runs:
        for c in r.text:
            cp = ord(c)
            if cp == 0x200B: hid['ZWSP'] += 1
            elif cp == 0x00AD: hid['SoftHyphen'] += 1
            elif cp == 0xFEFF: hid['BOM'] += 1
w(f"[Hidden Chars] {'PASS' if not hid else 'WARN'}: {dict(hid) if hid else 'none'}")

# 8. Style distribution
styles = defaultdict(int)
for p in doc.paragraphs:
    styles[p.style.name if p.style else 'None'] += 1
w(f"\n[Styles]")
for s, c in sorted(styles.items(), key=lambda x: -x[1])[:15]:
    w(f"  {s}: {c}")

out.close()
print("Done - audit_p1_results.txt")
