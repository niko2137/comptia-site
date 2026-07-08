"""Parts 3,4,5,6: Images, Content, Q&A, Modules"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from collections import Counter

print("Loading KDP-final.docx...")
doc = Document(r'public\reference\book\KDP-final.docx')
print(f"Loaded: {len(doc.paragraphs)} paragraphs")
out = open('audit_p345_results.txt', 'w', encoding='utf-8')

def w(t=""):
    out.write(t + "\n")

# === PART 3: IMAGES ===
w("PART 3: IMAGES")
w("=" * 60)

imgs = []
for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        centered = para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        for d in drawings:
            docPr = d.find('.//' + qn('wp:docPr'))
            name = docPr.get('name','') if docPr is not None else ''
            alt = docPr.get('descr','') if docPr is not None else ''
            prev_t = doc.paragraphs[i-1].text[:40] if i > 0 else ''
            imgs.append({'i':i,'name':name,'alt':alt,'centered':centered,'prev':prev_t})

w(f"Total images: {len(imgs)}")
no_alt = [x for x in imgs if not x['alt']]
no_center = [x for x in imgs if not x['centered']]
w(f"Missing alt text: {len(no_alt)}")
for x in no_alt:
    w(f"  Para {x['i']}: name='{x['name']}' near: {x['prev']}")
w(f"Not centered: {len(no_center)}")
for x in no_center:
    w(f"  Para {x['i']}: name='{x['name']}' near: {x['prev']}")
w(f"\nImage inventory:")
for x in imgs:
    c = 'Y' if x['centered'] else 'N'
    a = 'Y' if x['alt'] else 'N'
    w(f"  [{c}/{a}] Para {x['i']:4d} | {x['name'][:35]:35s} | {x['alt'][:40]}")

# === PART 4: CONTENT ACCURACY ===
w("\n\nPART 4: COMPTIA A+ CONTENT ACCURACY")
w("=" * 60)

all_text = '\n'.join(p.text for p in doc.paragraphs)
all_lower = all_text.lower()

facts = [
    ('220-1201', 'Core 1 exam number'),
    ('220-1202', 'Core 2 exam number'),
    ('90 minutes', 'Exam time limit'),
    ('90 questions', 'Max questions per exam'),
    ('675', 'Core 1 passing score'),
    ('700', 'Core 2 passing score'),
    ('WPA3', 'WPA3 security'),
    ('Wi-Fi 6E', 'Wi-Fi 6E coverage'),
    ('Wi-Fi 7', 'Wi-Fi 7 (802.11be) - new for v15'),
    ('DDR5', 'DDR5 RAM'),
    ('PCIe 4.0', 'PCIe 4.0'),
    ('PCIe 5.0', 'PCIe 5.0 - new for v15'),
    ('USB4', 'USB4 standard'),
    ('TPM 2.0', 'TPM 2.0'),
    ('Windows 11', 'Windows 11'),
    ('Windows 10', 'Windows 10'),
    ('DORA', 'DHCP DORA'),
    ('IETAID', 'IETAID mnemonic'),
    ('Thread', 'Thread protocol (IoT)'),
    ('Matter', 'Matter protocol (IoT) - new'),
    ('Zero Trust', 'Zero Trust - new for v15'),
    ('SASE', 'SASE - new for v15'),
    ('MDR', 'Managed Detection and Response'),
]

w("\nCritical Facts Check:")
missing = []
for search, desc in facts:
    count = all_lower.count(search.lower())
    status = 'FOUND' if count > 0 else 'MISSING'
    w(f"  [{status}] {desc}: '{search}' ({count}x)")
    if count == 0:
        missing.append(desc)

w(f"\nMissing: {len(missing)}")
for m in missing:
    w(f"  - {m}")

# Port numbers verification
w("\nPort Number Accuracy:")
port_checks = [
    (20, 'FTP data'), (21, 'FTP control'), (22, 'SSH/SFTP'),
    (23, 'Telnet'), (25, 'SMTP'), (53, 'DNS'),
    (67, 'DHCP server'), (68, 'DHCP client'), (80, 'HTTP'),
    (110, 'POP3'), (143, 'IMAP'), (161, 'SNMP'),
    (389, 'LDAP'), (443, 'HTTPS'), (445, 'SMB'),
    (587, 'SMTP TLS'), (636, 'LDAPS'), (993, 'IMAPS'),
    (995, 'POP3S'), (3389, 'RDP'), (5900, 'VNC'),
]
for port, proto in port_checks:
    if str(port) in all_text:
        w(f"  [OK] Port {port} ({proto})")
    else:
        w(f"  [MISSING] Port {port} ({proto})")

# Outdated info check
w("\nOutdated Information Check:")
outdated = [
    ('220-1001', 'Old Core 1 (v10)'),
    ('220-1002', 'Old Core 2 (v10)'),
    ('220-1101', 'Previous Core 1 (v14)'),
    ('220-1102', 'Previous Core 2 (v14)'),
]
for search, desc in outdated:
    count = all_text.count(search)
    if count:
        w(f"  [WARN] '{search}' found {count}x - {desc}")
    else:
        w(f"  [OK] No '{search}'")

# === PART 5: QUESTIONS & ANSWERS ===
w("\n\nPART 5: PRACTICE QUESTIONS & ANSWERS")
w("=" * 60)

questions = []
answers = []
in_q = False
in_a = False

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else ''
    
    if 'Practice Test' in text and 'Heading' in style:
        in_q = True
        in_a = False
        continue
    if ('Answer' in text and 'Heading' in style) or ('Scoring' in text and 'Heading' in style):
        in_q = False
        in_a = True
        continue
    if 'Heading 1' in style and (in_q or in_a):
        if 'Answer' not in text and 'Scoring' not in text and 'Practice' not in text:
            in_q = False
            in_a = False
            continue
    
    if in_q:
        m = re.match(r'^(\d+)\.\s+(.+)', text)
        if m:
            questions.append((int(m.group(1)), m.group(2)[:50], i))
    
    if in_a:
        m = re.match(r'^(\d+)\.\s+([A-D])\b', text)
        if m:
            answers.append((int(m.group(1)), m.group(2), i))

w(f"Questions found: {len(questions)}")
w(f"Answers found: {len(answers)}")

if questions:
    q_nums = [q[0] for q in questions]
    w(f"Question range: {min(q_nums)} to {max(q_nums)}")
    expected = set(range(1, max(q_nums)+1))
    missing_q = sorted(expected - set(q_nums))
    if missing_q:
        w(f"  [ISSUE] Missing question numbers: {missing_q[:20]}")
    else:
        w(f"  [OK] All question numbers sequential")
    
    dupes = [n for n, c in Counter(q_nums).items() if c > 1]
    if dupes:
        w(f"  [ISSUE] Duplicate questions: {dupes}")
    else:
        w(f"  [OK] No duplicate question numbers")

if answers:
    a_nums = [a[0] for a in answers]
    w(f"Answer range: {min(a_nums)} to {max(a_nums)}")
    invalid = [a for a in answers if a[1] not in 'ABCDE']
    if invalid:
        w(f"  [ISSUE] Invalid answer values: {invalid[:10]}")
    else:
        w(f"  [OK] All answers valid (A-D)")

if questions and answers:
    if len(questions) == len(answers):
        w(f"  [OK] Q/A count match: {len(questions)} each")
    else:
        w(f"  [ISSUE] Q/A MISMATCH: {len(questions)} questions vs {len(answers)} answers")
    
    # Check number alignment
    mismatched = []
    for q, a in zip(questions, answers):
        if q[0] != a[0]:
            mismatched.append((q[0], a[0]))
    if mismatched:
        w(f"  [ISSUE] Number misalignment: {mismatched[:10]}")
    else:
        w(f"  [OK] All Q/A numbers align")

# === PART 6: MODULE ORDER ===
w("\n\nPART 6: MODULE ORDER & OBJECTIVES")
w("=" * 60)

modules = []
for i, para in enumerate(doc.paragraphs):
    if para.style and para.style.name == 'Heading 1':
        modules.append((i, para.text.strip()))

w(f"Total H1 entries: {len(modules)}")
w("\nModule sequence:")
section = 'INTRO'
for idx, (pi, title) in enumerate(modules):
    if 'CORE 1' in title.upper():
        section = 'CORE 1'
        w(f"\n  ═══ {title} ═══")
    elif 'CORE 2' in title.upper():
        section = 'CORE 2'
        w(f"\n  ═══ {title} ═══")
    elif 'EXAM PREP' in title.upper() or 'STUDY TOOL' in title.upper():
        section = 'STUDY'
        w(f"\n  ═══ {title} ═══")
    elif 'Glossary' in title:
        section = 'GLOSSARY'
        w(f"\n  ═══ {title} ═══")
    else:
        w(f"  {idx+1}. [{section}] {title}")

# Check objectives exist after module headings
w("\nModule Objectives Check:")
missing_obj = []
for pi, title in modules:
    if 'Module' not in title:
        continue
    found = False
    for j in range(pi+1, min(pi+8, len(doc.paragraphs))):
        if 'Objective' in doc.paragraphs[j].text:
            found = True
            break
    if not found:
        missing_obj.append(title[:50])

if missing_obj:
    w(f"  [WARN] {len(missing_obj)} modules without objectives:")
    for m in missing_obj:
        w(f"    - {m}")
else:
    w(f"  [OK] All modules have objectives section")

# Required topics coverage
w("\nTopic Coverage:")
topics = [
    'Mobile Devices','Laptop','Networking','Virtualization','Cloud',
    'Hardware','Display','Storage','Printer','Cable','Connector',
    'Operating System','Windows','Linux','macOS','Security','Malware',
    'Troubleshooting','IoT','Embedded','Scripting','Backup','Recovery',
    'RAID','Remote Access','MDM','Data Destruction','Licensing',
    'Artificial Intelligence','Change Management','Ticket',
]
missing_topics = [t for t in topics if t.lower() not in all_lower]
if missing_topics:
    w(f"  [WARN] Missing topics: {missing_topics}")
else:
    w(f"  [OK] All {len(topics)} required topics covered")

# Callout audit
w("\nCallout Audit:")
trap_c = exam_c = pro_c = 0
trap_bad = exam_bad = pro_bad = 0
for para in doc.paragraphs:
    if not para.style: continue
    s = para.style.name
    t = para.text.strip()
    if s == 'Tip -Trap':
        trap_c += 1
        if not t.startswith('\u26A0'): trap_bad += 1
    elif s == 'Tip -Exam':
        exam_c += 1
        if not t.startswith('\u2139'): exam_bad += 1
    elif s == 'Tip -Pro':
        pro_c += 1
        if not t.startswith('\u2714'): pro_bad += 1

w(f"  Exam Trap: {trap_c} total, {trap_bad} non-standard")
w(f"  Exam Tip: {exam_c} total, {exam_bad} non-standard")
w(f"  Pro Tip: {pro_c} total, {pro_bad} non-standard")

# Glossary bullets
w("\nGlossary Bullets:")
colors = {}
for para in doc.paragraphs:
    for run in para.runs:
        if '\u25cf' in run.text:
            try:
                c = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'none'
            except:
                c = 'err'
            colors[c] = colors.get(c, 0) + 1
for c, n in colors.items():
    w(f"  {c}: {n}")

out.close()
print("Done - audit_p345_results.txt")
