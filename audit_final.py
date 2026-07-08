"""
FINAL AUDIT SCRIPT for KDP-final.docx
Checks for:
1. Remaining tables (should be 0)
2. Remaining emoji in headings
3. Remaining emoji in body text (excluding allowed: checkmark, info, warning)
4. Page breaks
5. Headers/footers with content
6. Images without alt text
7. Images not centered
8. Callout consistency
9. Glossary bullet colors
10. Duplicate/repeated text patterns
11. Orphaned formatting (empty runs, stray characters)
12. Style inconsistencies
13. TOC remnants
14. Broken bold patterns (** markers left as text)
15. Unicode issues (variation selectors, ZWJ outside allowed)
16. Excessive whitespace in paragraphs
"""

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
import re

SOURCE = r'public\reference\book\KDP-final.docx'
REPORT = 'final_audit_report.txt'

# Allowed emoji
ALLOWED = {'\u2714', '\u2139', '\u26A0'}  # ✔, ℹ, ⚠

def is_emoji(char):
    cp = ord(char)
    if char in ALLOWED:
        return False
    if 0x1F000 <= cp <= 0x1FFFF: return True
    if 0x2600 <= cp <= 0x26FF and char not in ALLOWED: return True
    if 0x2700 <= cp <= 0x27BF and char != '\u2714': return True
    if 0x2300 <= cp <= 0x23FF: return True
    if 0x2B00 <= cp <= 0x2BFF: return True
    if 0xFE00 <= cp <= 0xFE0F: return True  # Variation selectors
    if cp == 0x200D: return True  # ZWJ
    return False

def audit():
    out = open(REPORT, 'w', encoding='utf-8')
    def p(text=""):
        out.write(text + "\n")
    
    p("=" * 70)
    p("FINAL AUDIT - KDP-final.docx")
    p("=" * 70)
    
    doc = Document(SOURCE)
    p(f"\nDocument: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables, {len(doc.sections)} sections")
    
    issues = []
    warnings = []
    
    # === 1. TABLES ===
    p("\n--- 1. TABLES ---")
    if doc.tables:
        issues.append(f"{len(doc.tables)} tables remain")
        p(f"  ISSUE: {len(doc.tables)} tables still in document")
        for i, t in enumerate(doc.tables[:5]):
            if t.rows:
                cells = t.rows[0].cells
                header = " | ".join(c.text.strip()[:30] for c in cells[:4])
                p(f"    Table {i+1}: {header}")
    else:
        p("  OK: Zero tables")
    
    # === 2. EMOJI IN HEADINGS ===
    p("\n--- 2. EMOJI IN HEADINGS ---")
    heading_emoji = []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            for char in para.text:
                if is_emoji(char):
                    heading_emoji.append((para.style.name, para.text[:60], hex(ord(char))))
                    break
    
    if heading_emoji:
        issues.append(f"{len(heading_emoji)} headings still have emoji")
        p(f"  ISSUE: {len(heading_emoji)} headings with emoji")
        for style, text, cp in heading_emoji[:10]:
            p(f"    [{style}] {text} (char: {cp})")
    else:
        p("  OK: No emoji in headings")
    
    # === 3. EMOJI IN BODY ===
    p("\n--- 3. EMOJI IN BODY TEXT ---")
    body_emoji = []
    skip_styles = {'Glossary High', 'Glossary Medium', 'Glossary Low', 'Glossary Key'}
    for para in doc.paragraphs:
        if not para.style or para.style.name.startswith('Heading'):
            continue
        if para.style.name in skip_styles:
            continue
        # Skip callout styles
        if 'Tip' in para.style.name:
            continue
        
        for char in para.text:
            if is_emoji(char):
                # Check if it's a variation selector after allowed emoji
                idx = para.text.index(char)
                if idx > 0 and para.text[idx-1] in ALLOWED and 0xFE00 <= ord(char) <= 0xFE0F:
                    continue
                body_emoji.append((para.style.name, para.text[:80], hex(ord(char))))
                break
    
    if body_emoji:
        # Separate true issues from minor ones
        real_emoji = [b for b in body_emoji if int(b[2], 16) >= 0x1F000]
        variation_only = [b for b in body_emoji if 0xFE00 <= int(b[2], 16) <= 0xFE0F]
        zwj_only = [b for b in body_emoji if int(b[2], 16) == 0x200D]
        other = [b for b in body_emoji if b not in real_emoji and b not in variation_only and b not in zwj_only]
        
        if real_emoji:
            issues.append(f"{len(real_emoji)} body paragraphs have real emoji")
            p(f"  ISSUE: {len(real_emoji)} body paragraphs with emoji (U+1Fxxx)")
            for style, text, cp in real_emoji[:10]:
                p(f"    [{style}] {text}")
        if variation_only:
            warnings.append(f"{len(variation_only)} stray variation selectors")
            p(f"  WARN: {len(variation_only)} stray variation selectors")
        if zwj_only:
            warnings.append(f"{len(zwj_only)} stray ZWJ characters")
            p(f"  WARN: {len(zwj_only)} stray ZWJ")
        if other:
            p(f"  INFO: {len(other)} paragraphs with misc symbols (U+2600-27FF range)")
            for style, text, cp in other[:5]:
                p(f"    [{style}] {cp}: {text}")
        if not real_emoji and not variation_only and not zwj_only:
            p("  OK: No problematic emoji in body")
    else:
        p("  OK: No emoji in body text")
    
    # === 4. PAGE BREAKS ===
    p("\n--- 4. PAGE BREAKS ---")
    page_breaks = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            if pPr.find(qn('w:pageBreakBefore')) is not None:
                page_breaks += 1
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    page_breaks += 1
    
    if page_breaks:
        issues.append(f"{page_breaks} page breaks remain")
        p(f"  ISSUE: {page_breaks} page breaks")
    else:
        p("  OK: No page breaks")
    
    # === 5. HEADERS/FOOTERS ===
    p("\n--- 5. HEADERS & FOOTERS ---")
    hf_issues = 0
    for i, section in enumerate(doc.sections):
        for para in section.header.paragraphs:
            if para.text.strip():
                hf_issues += 1
                p(f"  ISSUE: Section {i+1} header: '{para.text.strip()[:50]}'")
        for para in section.footer.paragraphs:
            if para.text.strip():
                hf_issues += 1
                p(f"  ISSUE: Section {i+1} footer: '{para.text.strip()[:50]}'")
    
    if hf_issues:
        issues.append(f"{hf_issues} headers/footers with content")
    else:
        p("  OK: All headers/footers empty")
    
    # === 6. IMAGES ===
    p("\n--- 6. IMAGES ---")
    img_total = 0
    img_no_alt = 0
    img_not_centered = 0
    
    for para in doc.paragraphs:
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            img_total += 1
            if para.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                img_not_centered += 1
            for d in drawings:
                docPr = d.find('.//' + qn('wp:docPr'))
                if docPr is not None:
                    if not docPr.get('descr', ''):
                        img_no_alt += 1
                else:
                    img_no_alt += 1
    
    p(f"  Total image paragraphs: {img_total}")
    if img_no_alt:
        issues.append(f"{img_no_alt} images without alt text")
        p(f"  ISSUE: {img_no_alt} images missing alt text")
    else:
        p("  OK: All images have alt text")
    if img_not_centered:
        warnings.append(f"{img_not_centered} images not centered")
        p(f"  WARN: {img_not_centered} images not centered")
    else:
        p("  OK: All images centered")
    
    # === 7. CALLOUTS ===
    p("\n--- 7. CALLOUTS ---")
    callout_issues = []
    for para in doc.paragraphs:
        if not para.style:
            continue
        style = para.style.name
        text = para.text.strip()
        
        if style == 'Tip -Trap':
            if not text.startswith('\u26A0'):
                callout_issues.append(('Trap missing ⚠️', text[:60]))
        elif style == 'Tip -Exam':
            if not text.startswith('\u2139'):
                callout_issues.append(('Exam missing ℹ️', text[:60]))
        elif style == 'Tip -Pro':
            if not text.startswith('\u2714'):
                callout_issues.append(('Pro missing ✔', text[:60]))
    
    if callout_issues:
        issues.append(f"{len(callout_issues)} callouts not standardized")
        p(f"  ISSUE: {len(callout_issues)} callouts")
        for issue_type, text in callout_issues[:5]:
            p(f"    {issue_type}: {text}")
    else:
        p("  OK: All callouts standardized")
    
    # === 8. GLOSSARY BULLETS ===
    p("\n--- 8. GLOSSARY BULLETS ---")
    bullet_colors = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if '\u25cf' in run.text:
                try:
                    color = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'no-color'
                except:
                    color = 'error'
                bullet_colors[color] = bullet_colors.get(color, 0) + 1
    
    if bullet_colors:
        for color, count in bullet_colors.items():
            p(f"  {color}: {count}")
        expected = {'FF0000', 'FFC000', '00B300'}
        if expected.issubset(set(bullet_colors.keys())):
            p("  OK: All 3 priority colors present")
        else:
            missing = expected - set(bullet_colors.keys())
            warnings.append(f"Missing glossary colors: {missing}")
            p(f"  WARN: Missing colors: {missing}")
    else:
        issues.append("No glossary bullets found")
        p("  ISSUE: No glossary bullets found!")
    
    # === 9. TOC REMNANTS ===
    p("\n--- 9. TOC REMNANTS ---")
    toc_remnants = 0
    for para in doc.paragraphs:
        if para.style and para.style.name.lower().startswith('toc'):
            toc_remnants += 1
    
    if toc_remnants:
        issues.append(f"{toc_remnants} TOC-styled paragraphs remain")
        p(f"  ISSUE: {toc_remnants} paragraphs with TOC styles")
    else:
        p("  OK: No TOC remnants")
    
    # === 10. ASTERISK MARKERS (** left as text) ===
    p("\n--- 10. LITERAL ** MARKERS ---")
    asterisk_issues = []
    for para in doc.paragraphs:
        if '**' in para.text:
            asterisk_issues.append((para.style.name if para.style else 'None', para.text[:80]))
    
    if asterisk_issues:
        issues.append(f"{len(asterisk_issues)} paragraphs with literal ** markers")
        p(f"  ISSUE: {len(asterisk_issues)} paragraphs with ** (bold not applied)")
        for style, text in asterisk_issues[:10]:
            p(f"    [{style}] {text}")
    else:
        p("  OK: No literal ** markers")
    
    # === 11. DUPLICATE CONSECUTIVE TEXT ===
    p("\n--- 11. DUPLICATE/REPEATED TEXT ---")
    prev_text = ""
    duplicates = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and text == prev_text and len(text) > 20:
            duplicates.append(text[:60])
        prev_text = text
    
    if duplicates:
        warnings.append(f"{len(duplicates)} consecutive duplicate paragraphs")
        p(f"  WARN: {len(duplicates)} consecutive duplicates")
        for d in duplicates[:5]:
            p(f"    '{d}'")
    else:
        p("  OK: No consecutive duplicates")
    
    # === 12. REPEATED TEXT WITHIN PARAGRAPHS ===
    p("\n--- 12. INTERNAL REPETITION (text repeated 3+ times in same para) ---")
    internal_repeats = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) > 50:
            # Check if text has obvious repetition (same phrase 3+ times)
            words = text.split()
            if len(words) >= 9:
                # Check first 3 words pattern
                chunk = ' '.join(words[:3])
                if text.count(chunk) >= 3:
                    internal_repeats.append((para.style.name if para.style else 'None', text[:80]))
    
    if internal_repeats:
        issues.append(f"{len(internal_repeats)} paragraphs with internal repetition")
        p(f"  ISSUE: {len(internal_repeats)} paragraphs repeat same text internally")
        for style, text in internal_repeats[:10]:
            p(f"    [{style}] {text}")
    else:
        p("  OK: No internal text repetition")
    
    # === 13. EMPTY PARAGRAPHS WITH STYLES ===
    p("\n--- 13. STYLE DISTRIBUTION ---")
    styles = {}
    empty_styled = 0
    for para in doc.paragraphs:
        style = para.style.name if para.style else 'None'
        styles[style] = styles.get(style, 0) + 1
        if not para.text.strip() and style not in ('Normal', 'Body Text', 'First Paragraph'):
            # Empty paragraph with a non-default style
            has_img = bool(para._element.findall('.//' + qn('w:drawing')))
            if not has_img:
                empty_styled += 1
    
    p("  Style counts:")
    for style, count in sorted(styles.items(), key=lambda x: -x[1])[:15]:
        p(f"    {style}: {count}")
    
    if empty_styled > 10:
        warnings.append(f"{empty_styled} empty paragraphs with non-default styles")
        p(f"\n  WARN: {empty_styled} empty paragraphs with non-default styles")
    
    # === 14. BULLET CHARACTER CHECK ===
    p("\n--- 14. BULLET/LIST FORMATTING ---")
    bullet_chars = {'\u2022': 0, '\u25cf': 0, '\u2013': 0, '\u2014': 0}
    for para in doc.paragraphs:
        for char, _ in bullet_chars.items():
            if char in para.text:
                bullet_chars[char] += para.text.count(char)
    
    p(f"  Bullet (•): {bullet_chars[chr(0x2022)]}")
    p(f"  Filled circle (●): {bullet_chars[chr(0x25cf)]}")
    p(f"  En dash (–): {bullet_chars[chr(0x2013)]}")
    p(f"  Em dash (—): {bullet_chars[chr(0x2014)]}")
    
    # === 15. UNICODE ISSUES ===
    p("\n--- 15. UNICODE / INVISIBLE CHARACTERS ---")
    unicode_issues = []
    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            for i, char in enumerate(text):
                cp = ord(char)
                # Check for problematic invisible chars
                if cp == 0x200B:  # Zero-width space
                    unicode_issues.append(('ZWSP', para.text[:50]))
                    break
                elif cp == 0x00AD:  # Soft hyphen
                    unicode_issues.append(('Soft Hyphen', para.text[:50]))
                    break
                elif cp == 0xFEFF:  # BOM
                    unicode_issues.append(('BOM', para.text[:50]))
                    break
                elif cp == 0x00A0 and i > 0 and i < len(text) - 1:
                    # Non-breaking space in middle of text (might be intentional)
                    pass
    
    if unicode_issues:
        warnings.append(f"{len(unicode_issues)} invisible character issues")
        p(f"  WARN: {len(unicode_issues)} invisible characters found")
        seen_types = set()
        for utype, text in unicode_issues[:10]:
            if utype not in seen_types:
                p(f"    {utype}: {text}")
                seen_types.add(utype)
    else:
        p("  OK: No problematic invisible characters")
    
    # === 16. WHITESPACE ISSUES ===
    p("\n--- 16. WHITESPACE ISSUES ---")
    ws_issues = 0
    leading_space = 0
    for para in doc.paragraphs:
        text = para.text
        if text and text != text.strip():
            if text.startswith(' ') or text.startswith('\t'):
                leading_space += 1
        if '  ' in text:  # Double spaces
            ws_issues += 1
    
    if ws_issues > 20:
        warnings.append(f"{ws_issues} paragraphs with double spaces")
        p(f"  WARN: {ws_issues} paragraphs with double spaces")
    else:
        p(f"  OK: {ws_issues} paragraphs with double spaces (minimal)")
    if leading_space > 20:
        warnings.append(f"{leading_space} paragraphs with leading whitespace")
        p(f"  WARN: {leading_space} paragraphs with leading spaces/tabs")
    else:
        p(f"  OK: {leading_space} paragraphs with leading whitespace (minimal)")
    
    # === SUMMARY ===
    p("\n" + "=" * 70)
    p("AUDIT SUMMARY")
    p("=" * 70)
    p(f"\n  ISSUES (must fix): {len(issues)}")
    for issue in issues:
        p(f"    ❌ {issue}")
    
    p(f"\n  WARNINGS (review): {len(warnings)}")
    for w in warnings:
        p(f"    ⚠️ {w}")
    
    if not issues and not warnings:
        p("\n  ✅ DOCUMENT APPEARS CLEAN!")
    elif not issues:
        p("\n  ✅ No critical issues. Warnings are informational.")
    
    out.close()
    print(f"Audit complete - see {REPORT}")

if __name__ == '__main__':
    audit()
