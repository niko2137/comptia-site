"""
KDP CONVERSION SCRIPT - Version 3
Based on KDP_CONVERSION_PROMPT_v3.md

Converts comptia-book.docx to KDP-ready reflowable ebook format.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import re
import shutil
import os
import copy

# Configuration
SOURCE_FILE = r'public\reference\book\comptia-book.docx'
OUTPUT_FILE = r'public\reference\book\comptia-book-KDP-v3.docx'
BACKUP_FILE = r'public\reference\book\comptia-book-BACKUP-v3.docx'

# === EXCEPTION LISTS (identified by header row content) ===

# 3A: Comparison tables -> bullet lists
# Identified by matching header text patterns
COMPARISON_TABLES = [
    # Wi-Fi 6E (table 15)
    lambda h: 'Wi-Fi 6' in h and '6E' in h,
    # TCP vs UDP (table 27)
    lambda h: 'TCP' in h and 'UDP' in h and 'Feature' in h,
    # OSI vs TCP/IP side-by-side (table 29)
    lambda h: 'OSI Layer' in h and 'TCP/IP' in h,
    # Hypervisor Types (table 41)
    lambda h: 'Type 1 Hypervisor' in h or 'Type 2 Hypervisor' in h,
    # Snapshots and Cloning (table 43)
    lambda h: 'Snapshot' in h and 'Clone' in h,
    # Storage Drive Comparison (table 70)
    lambda h: 'HDD' in h and 'SSD' in h and 'Feature' in h,
    # MBR vs GPT (table 72)
    lambda h: 'MBR' in h and 'GPT' in h,
    # File System Comparison (table 75)
    lambda h: 'NTFS' in h and 'FAT32' in h,
    # What is an Embedded System? (table 91)
    lambda h: 'Embedded System' in h and 'General-Purpose' in h,
    # Thread vs Other IoT Protocols (table 94)
    lambda h: 'Protocol' in h and 'IP-Based' in h and 'Range' in h,
    # Windows 10/11 Edition Comparison (table 96)
    lambda h: 'Home' in h and 'Pro' in h and 'Enterprise' in h,
    # 32-bit vs 64-bit (table 100)
    lambda h: '32-bit' in h and '64-bit' in h,
]


# 3B: Special formatting tables (sequential process/procedure)
# Shape A: 2-col with "Step | Action" where step has embedded label
# Shape B: 3-col with "Step | Official Action | Deep Dive"
SPECIAL_TABLES = [
    # TCP/IP Model (table 26) - actually 4-col reference, treat as 3A bullet
    lambda h: h.startswith('Layer') and 'Name' in h and 'Function' in h and 'Example' in h,
    # OSI Model 7 layers (table 28) - 4-col reference, treat as 3A bullet
    lambda h: h.startswith('Layer') and 'Name' in h and 'Function' in h and 'Key Protocols' in h,
    # Motherboard Components (table 50) - special multirow with merged cells
    lambda h: '#' in h and 'Component' in h and 'Exam Notes' in h and 'Priority' in h,
    # 7-Step Malware Removal (table 109) - Shape B
    lambda h: 'Step' in h and 'Action' in h and 'Exam Focus' in h,
    # 6-Step Troubleshooting Method (table 142) - Shape B
    lambda h: 'Step' in h and 'Official Action' in h and 'Deep Dive' in h,
    # Ticket Lifecycle (table 146) - 2-col Step|Details
    lambda h: 'Step' in h and 'Details' in h and h.count('|') == 1,
    # IETAID (table 177) - Shape A (Step | Action)
    lambda h: 'Step' in h and 'Action' in h and h.count('|') == 1 and 'Exam' not in h,
]

# Statistics tracking
stats = {
    'tables_flattened': 0,
    'tables_to_bullets': 0,
    'tables_special': 0,
    'headings_emoji_stripped': 0,
    'body_emoji_stripped': 0,
    'toc_removed': 0,
    'callouts_fixed': 0,
    'page_breaks_removed': 0,
    'images_centered': 0,
    'images_missing_alt': 0,
    'headers_cleared': 0,
    'footers_cleared': 0,
    'blank_paragraphs_removed': 0,
    'glossary_bullets_preserved': 0,
    'checkmarks_replaced': 0,
}


# Emoji pattern - BROAD coverage
# Allowed emoji that should NOT be removed: ✔ (U+2714), ℹ️ (U+2139+FE0F), ⚠️ (U+26A0+FE0F)
ALLOWED_EMOJI = {'\u2714', '\u2139', '\u26A0'}

def is_allowed_emoji(char):
    """Check if a character is in the allowed emoji list."""
    return char in ALLOWED_EMOJI

def strip_emoji(text, keep_allowed=True):
    """Remove emoji from text. If keep_allowed=True, preserves ✔, ℹ️, ⚠️."""
    result = []
    i = 0
    while i < len(text):
        char = text[i]
        cp = ord(char)
        
        # Check if it's an allowed emoji
        if keep_allowed and char in ALLOWED_EMOJI:
            result.append(char)
            i += 1
            continue
        
        # Variation selectors (U+FE00-FE0F) - remove unless following allowed
        if 0xFE00 <= cp <= 0xFE0F:
            # Keep if previous char is allowed emoji
            if result and result[-1] in ALLOWED_EMOJI:
                result.append(char)
            # else skip
            i += 1
            continue
        
        # Zero-width joiner
        if cp == 0x200D:
            i += 1
            continue
        
        # Emoji ranges to remove
        is_emoji = False
        if 0x1F000 <= cp <= 0x1FFFF:  # Emoticons, symbols, etc
            is_emoji = True
        elif 0x2600 <= cp <= 0x26FF:  # Misc symbols
            if char not in ALLOWED_EMOJI:
                is_emoji = True
        elif 0x2700 <= cp <= 0x27BF:  # Dingbats
            if char != '\u2714':  # Keep checkmark
                is_emoji = True
        elif 0x2300 <= cp <= 0x23FF:  # Misc technical
            is_emoji = True
        elif 0x2B00 <= cp <= 0x2BFF:  # Misc symbols/arrows
            is_emoji = True
        elif 0x200D == cp:  # ZWJ
            is_emoji = True
        elif 0xE0020 <= cp <= 0xE007F:  # Tags
            is_emoji = True
        elif 0x1F900 <= cp <= 0x1F9FF:  # Supplemental symbols
            is_emoji = True
        
        if is_emoji:
            i += 1
            continue
        
        result.append(char)
        i += 1
    
    return ''.join(result)


def deduplicate_cells(row):
    """Deduplicate merged cells by element identity."""
    seen = set()
    unique = []
    for cell in row.cells:
        cell_id = id(cell._element)
        if cell_id not in seen:
            seen.add(cell_id)
            unique.append(cell)
    return unique


def get_table_header_text(table):
    """Get the header row as pipe-separated text for identification."""
    if len(table.rows) == 0:
        return ""
    cells = deduplicate_cells(table.rows[0])
    return " | ".join(c.text.strip()[:50] for c in cells)


def is_comparison_table(header_text):
    """Check if table matches any comparison table pattern."""
    for matcher in COMPARISON_TABLES:
        if matcher(header_text):
            return True
    return False


def is_special_table(header_text):
    """Check if table matches any special formatting pattern."""
    for matcher in SPECIAL_TABLES:
        if matcher(header_text):
            return True
    return False


def get_special_table_type(header_text, table):
    """Determine if special table is Shape A, B, or reference-style."""
    cells = deduplicate_cells(table.rows[0])
    num_cols = len(cells)
    header_lower = header_text.lower()
    
    # TCP/IP Model and OSI Model are really reference tables -> bullet format
    if 'layer' in header_lower and 'name' in header_lower and 'function' in header_lower:
        return 'bullet'  # Treat like comparison/bullet table
    
    # Motherboard Components - special merged cell table
    if '#' in header_text and 'component' in header_lower:
        return 'motherboard'
    
    # Shape A: 2-col "Step | Action" with embedded labels
    if num_cols == 2 and 'step' in header_lower and 'action' in header_lower:
        # Check if step values have embedded labels like "1 - I"
        if len(table.rows) > 1:
            first_step = deduplicate_cells(table.rows[1])[0].text.strip()
            if ' - ' in first_step or first_step.startswith('1'):
                return 'shape_a'
    
    # Shape B: 3-col "Step | Title | Description"
    if num_cols == 3 and 'step' in header_lower:
        return 'shape_b'
    
    # Ticket lifecycle: 2-col "Step | Details" with numbered steps
    if num_cols == 2 and 'step' in header_lower and 'details' in header_lower:
        return 'shape_a'
    
    return 'shape_b'  # Default


def flatten_table_standard(table, is_first_after_heading=False):
    """Standard table flattening: 2-col, 3-col, 4+ col rules."""
    rows = table.rows
    if len(rows) == 0:
        return []
    
    header_cells = deduplicate_cells(rows[0])
    headers = [c.text.strip() for c in header_cells]
    num_cols = len(headers)
    
    paragraphs = []
    
    for row_idx in range(1, len(rows)):
        cells = deduplicate_cells(rows[row_idx])
        values = [c.text.strip() for c in cells]
        
        # Skip empty rows
        if not any(values):
            continue
        
        # Handle full-width merged rows (single cell spanning all columns)
        if len(values) == 1 and len(cells) == 1:
            # This is likely a section header or note row
            text = values[0]
            if text:
                paragraphs.append(('Body Text', f"**{text}**"))
            continue
        
        # Determine style
        style = 'First Paragraph' if row_idx == 1 and is_first_after_heading else 'Body Text'
        
        if num_cols == 2:
            # 2-column: Term - description
            term = values[0] if len(values) > 0 else ''
            desc = values[1] if len(values) > 1 else ''
            if term and desc:
                text = f"**{term}** \u2014 {desc}"
            elif term:
                text = f"**{term}**"
            else:
                text = desc
            paragraphs.append((style, text))
        
        elif num_cols == 3:
            # 3-column: Term + soft break + Header1: value + soft break + Header2: value
            term = values[0] if len(values) > 0 else ''
            parts = [f"**{term}**"] if term else []
            for i in range(1, min(len(values), num_cols)):
                if values[i]:
                    parts.append(f"{headers[i]}: {values[i]}")
            text = '\n'.join(parts)
            paragraphs.append((style, text))
        
        else:
            # 4+ columns: Term + soft break + HeaderN: valueN per column
            term = values[0] if len(values) > 0 else ''
            parts = [f"**{term}**"] if term else []
            for i in range(1, min(len(values), num_cols)):
                if values[i]:
                    header_label = headers[i] if i < len(headers) else f"Col{i}"
                    parts.append(f"{header_label}: {values[i]}")
            text = '\n'.join(parts)
            paragraphs.append((style, text))
    
    stats['tables_flattened'] += 1
    return paragraphs


def flatten_table_bullets(table, header_text):
    """Convert comparison table to bullet-list format."""
    rows = table.rows
    if len(rows) == 0:
        return []
    
    header_cells = deduplicate_cells(rows[0])
    headers = [c.text.strip() for c in header_cells]
    
    # Check if this is the Windows Edition Comparison (needs ✅->YES, ❌->NO)
    is_windows_edition = 'Home' in header_text and 'Pro' in header_text and 'Enterprise' in header_text
    
    paragraphs = []
    
    for row_idx in range(1, len(rows)):
        cells = deduplicate_cells(rows[row_idx])
        values = [c.text.strip() for c in cells]
        
        if not any(values):
            continue
        
        # Handle full-width merged rows
        if len(values) == 1:
            if values[0]:
                paragraphs.append(('Body Text', f"**{values[0]}**"))
            continue
        
        # First column = bold non-bulleted heading
        term = values[0] if values else ''
        if term:
            paragraphs.append(('Body Text', f"**{term}**"))
        
        # Remaining columns = bullet items
        for i in range(1, min(len(values), len(headers))):
            val = values[i]
            if is_windows_edition:
                val = val.replace('\u2705', 'YES').replace('\u274C', 'NO')
                val = val.replace('\u2714', 'YES')  # Also handle ✔
                stats['checkmarks_replaced'] += 1
            
            header_label = headers[i]
            if val:
                paragraphs.append(('Body Text', f"\u2022 {header_label}: {val}"))
    
    stats['tables_to_bullets'] += 1
    return paragraphs


def flatten_table_special(table, header_text):
    """Convert special process/procedure tables."""
    table_type = get_special_table_type(header_text, table)
    rows = table.rows
    
    if len(rows) == 0:
        return []
    
    header_cells = deduplicate_cells(rows[0])
    headers = [c.text.strip() for c in header_cells]
    
    paragraphs = []
    
    if table_type == 'bullet':
        # Treat like a comparison/bullet table
        return flatten_table_bullets(table, header_text)
    
    elif table_type == 'motherboard':
        # Motherboard Components - has merged section headers
        for row_idx in range(1, len(rows)):
            cells = deduplicate_cells(rows[row_idx])
            values = [c.text.strip() for c in cells]
            
            if not any(values):
                continue
            
            # Single-cell merged row = section header
            if len(values) == 1 or (len(values) >= 1 and not any(values[1:])):
                paragraphs.append(('Body Text', f"\n**{values[0]}**"))
                continue
            
            # Normal data row: # | Component | Exam Notes | Priority
            num = values[0] if len(values) > 0 else ''
            component = values[1] if len(values) > 1 else ''
            notes = values[2] if len(values) > 2 else ''
            priority = values[3] if len(values) > 3 else ''
            
            parts = []
            if num and component:
                parts.append(f"**{num}. {component}**")
            elif component:
                parts.append(f"**{component}**")
            if notes:
                parts.append(f"Exam Notes: {notes}")
            if priority:
                parts.append(f"Priority: {priority}")
            
            text = '\n'.join(parts) if len(parts) > 1 else (parts[0] if parts else '')
            if text:
                paragraphs.append(('Body Text', text))
    
    elif table_type == 'shape_a':
        # Shape A: "Step {Step value} — {Action value}"
        for row_idx in range(1, len(rows)):
            cells = deduplicate_cells(rows[row_idx])
            values = [c.text.strip() for c in cells]
            
            if not any(values):
                continue
            
            step_val = values[0] if len(values) > 0 else ''
            action_val = values[1] if len(values) > 1 else ''
            
            # Check if step already has "Step" prefix
            if step_val.lower().startswith('step'):
                text = f"**{step_val}** \u2014 {action_val}"
            else:
                text = f"**Step {step_val}** \u2014 {action_val}"
            paragraphs.append(('Body Text', text))
    
    elif table_type == 'shape_b':
        # Shape B: "Step {step} — {title}: {description}"
        for row_idx in range(1, len(rows)):
            cells = deduplicate_cells(rows[row_idx])
            values = [c.text.strip() for c in cells]
            
            if not any(values):
                continue
            
            step_val = values[0] if len(values) > 0 else ''
            title = values[1] if len(values) > 1 else ''
            desc = values[2] if len(values) > 2 else ''
            
            if step_val.lower().startswith('step'):
                text = f"**{step_val}** \u2014 {title}: {desc}" if desc else f"**{step_val}** \u2014 {title}"
            else:
                text = f"**Step {step_val}** \u2014 {title}: {desc}" if desc else f"**Step {step_val}** \u2014 {title}"
            paragraphs.append(('Body Text', text))
    
    stats['tables_special'] += 1
    return paragraphs


def create_paragraph_element(style_name, text):
    """Create a w:p element with styled text, handling bold markers and line breaks."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    p.append(pPr)
    
    # Split on newlines for soft breaks
    lines = text.split('\n')
    for line_idx, line in enumerate(lines):
        # Parse bold markers
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr.append(b)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = part[2:-2]
                r.append(t)
                p.append(r)
            elif part:
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = part
                r.append(t)
                p.append(r)
        
        # Add soft break between lines (not after last)
        if line_idx < len(lines) - 1:
            r = OxmlElement('w:r')
            br = OxmlElement('w:br')
            r.append(br)
            p.append(r)
    
    return p


def replace_table_with_paragraphs(table, paragraphs):
    """Replace a table element with text paragraphs."""
    table_element = table._element
    parent = table_element.getparent()
    table_position = list(parent).index(table_element)
    
    # Insert paragraphs before removing table
    for i, (style_name, text) in enumerate(paragraphs):
        p_elem = create_paragraph_element(style_name, text)
        parent.insert(table_position + i, p_elem)
    
    # Remove the table
    parent.remove(table_element)


def process_tables(doc):
    """Process all tables according to exception lists."""
    # Process in reverse to maintain element positions
    tables_to_process = []
    
    for i, table in enumerate(doc.tables):
        header_text = get_table_header_text(table)
        tables_to_process.append((i, table, header_text))
    
    for i, table, header_text in reversed(tables_to_process):
        if is_comparison_table(header_text):
            paragraphs = flatten_table_bullets(table, header_text)
            replace_table_with_paragraphs(table, paragraphs)
        elif is_special_table(header_text):
            paragraphs = flatten_table_special(table, header_text)
            replace_table_with_paragraphs(table, paragraphs)
        else:
            paragraphs = flatten_table_standard(table)
            replace_table_with_paragraphs(table, paragraphs)


def remove_toc(doc):
    """Remove TOC paragraphs (identified by toc 1 / toc 2 styles)."""
    to_remove = []
    for para in doc.paragraphs:
        if para.style and para.style.name.lower().startswith('toc'):
            to_remove.append(para._element)
    
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            stats['toc_removed'] += 1


def strip_emoji_from_headings(doc):
    """Remove emoji from Heading 1, 2, 3, 4, 5 paragraphs."""
    for para in doc.paragraphs:
        if not para.style or not para.style.name.startswith('Heading'):
            continue
        
        runs_to_remove = []
        for run in para.runs:
            original = run.text
            cleaned = strip_emoji(original, keep_allowed=False)
            # Strip leading/trailing whitespace that was around emoji
            cleaned = re.sub(r'^\s+', '', cleaned)
            
            if original != cleaned:
                stats['headings_emoji_stripped'] += 1
                run.text = cleaned
            
            # Mark empty runs for removal
            if not run.text.strip():
                runs_to_remove.append(run)
        
        # Remove empty runs
        for run in runs_to_remove:
            run._element.getparent().remove(run._element)


def strip_emoji_from_body(doc):
    """Remove emoji from Body text, Normal, and caption paragraphs.
    Preserves allowed emoji (checkmark, info, warning)."""
    target_styles = {'Normal', 'Body Text', 'First Paragraph', 'Compact',
                     'Captioned Figure', 'Image Caption', 'Caption1',
                     'ag-paragraph', 'Block Text'}
    
    for para in doc.paragraphs:
        if not para.style:
            continue
        style_name = para.style.name
        
        # Skip glossary styles - they have colored bullets
        if 'Glossary' in style_name or 'glossary' in style_name:
            continue
        # Skip callout styles - handled separately
        if 'Tip' in style_name:
            continue
        # Skip TOC (already removed but just in case)
        if style_name.lower().startswith('toc'):
            continue
        
        if style_name not in target_styles:
            continue
        
        for run in para.runs:
            # Don't touch glossary bullets
            if '\u25cf' in run.text:
                continue
            
            original = run.text
            cleaned = strip_emoji(original, keep_allowed=True)
            
            if original != cleaned:
                stats['body_emoji_stripped'] += 1
                run.text = cleaned


def fix_callouts(doc):
    """Standardize callout paragraphs: ⚠️ Exam Trap, ℹ️ Exam Tip, ✔ Pro Tip."""
    for para in doc.paragraphs:
        if not para.style:
            continue
        style_name = para.style.name
        
        if style_name == 'Tip -Trap':
            # Should start with ⚠️ Exam Trap:
            full_text = para.text
            if not full_text.startswith('\u26A0'):
                # Fix it - find first run with content
                if para.runs:
                    # Strip any existing emoji from first run
                    first_run = para.runs[0]
                    cleaned = strip_emoji(first_run.text, keep_allowed=False).lstrip()
                    if not cleaned.startswith('Exam Trap'):
                        first_run.text = '\u26A0\uFE0F Exam Trap: ' + cleaned
                    else:
                        first_run.text = '\u26A0\uFE0F ' + cleaned
                    stats['callouts_fixed'] += 1
        
        elif style_name == 'Tip -Exam':
            full_text = para.text
            if not full_text.startswith('\u2139'):
                if para.runs:
                    first_run = para.runs[0]
                    cleaned = strip_emoji(first_run.text, keep_allowed=False).lstrip()
                    if not cleaned.startswith('Exam Tip'):
                        first_run.text = '\u2139\uFE0F Exam Tip: ' + cleaned
                    else:
                        first_run.text = '\u2139\uFE0F ' + cleaned
                    stats['callouts_fixed'] += 1
        
        elif style_name == 'Tip -Pro':
            full_text = para.text
            # Replace 💡 with ✔
            if para.runs:
                first_run = para.runs[0]
                cleaned = strip_emoji(first_run.text, keep_allowed=False).lstrip()
                if not cleaned.startswith('Pro Tip'):
                    first_run.text = '\u2714 Pro Tip: ' + cleaned
                else:
                    first_run.text = '\u2714 ' + cleaned
                stats['callouts_fixed'] += 1


def remove_page_breaks(doc):
    """Remove all page breaks (pageBreakBefore and explicit w:br type=page)."""
    for para in doc.paragraphs:
        # Remove pageBreakBefore from paragraph properties
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            pb = pPr.find(qn('w:pageBreakBefore'))
            if pb is not None:
                pPr.remove(pb)
                stats['page_breaks_removed'] += 1
        
        # Remove explicit page break runs
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    run._element.remove(br)
                    stats['page_breaks_removed'] += 1


def center_images_and_check_alt(doc):
    """Center all images and verify alt text."""
    for para in doc.paragraphs:
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            # Center the paragraph
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            stats['images_centered'] += 1
            
            for drawing in drawings:
                # Check alt text
                docPr = drawing.find('.//' + qn('wp:docPr'))
                if docPr is not None:
                    alt = docPr.get('descr', '')
                    if not alt:
                        stats['images_missing_alt'] += 1
                        # Add placeholder alt text based on context
                        name = docPr.get('name', 'Image')
                        docPr.set('descr', f'{name} - CompTIA A+ Study Guide')
                else:
                    stats['images_missing_alt'] += 1


def clear_headers_footers(doc):
    """Remove all headers and footers content."""
    for section in doc.sections:
        # Clear headers
        for para in section.header.paragraphs:
            if para.text.strip():
                stats['headers_cleared'] += 1
            para.clear()
        
        # Clear footers
        for para in section.footer.paragraphs:
            if para.text.strip():
                stats['footers_cleared'] += 1
            para.clear()


def remove_excess_blanks(doc):
    """Collapse runs of 2+ empty paragraphs to max 1.
    IMPORTANT: A paragraph with images is NOT blank even if text is empty.
    Also keeps single spacer blanks adjacent to image paragraphs."""
    consecutive_empty = []
    to_remove = []
    prev_had_image = False
    
    for para in doc.paragraphs:
        # Preserve glossary bullets
        has_bullet = False
        for run in para.runs:
            if '\u25cf' in run.text:
                has_bullet = True
                try:
                    if run.font.color and run.font.color.rgb:
                        stats['glossary_bullets_preserved'] += 1
                except:
                    pass
        
        if has_bullet:
            # Process accumulated blanks before resetting
            if len(consecutive_empty) >= 2 and not prev_had_image:
                to_remove.extend(consecutive_empty[1:])
            consecutive_empty = []
            prev_had_image = False
            continue
        
        # Check if paragraph has images - if so, it's NOT blank
        has_images = bool(para._element.findall('.//' + qn('w:drawing')))
        if has_images:
            # Process accumulated blanks before image
            if len(consecutive_empty) >= 2 and not prev_had_image:
                to_remove.extend(consecutive_empty[1:])
            consecutive_empty = []
            prev_had_image = True
            continue
        
        if not para.text.strip():
            consecutive_empty.append(para._element)
        else:
            if len(consecutive_empty) >= 2 and not prev_had_image:
                to_remove.extend(consecutive_empty[1:])
            consecutive_empty = []
            prev_had_image = False
    
    # Handle trailing empties
    if len(consecutive_empty) >= 2 and not prev_had_image:
        to_remove.extend(consecutive_empty[1:])
    
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            stats['blank_paragraphs_removed'] += 1


def main():
    # Set up logging to file
    import io
    log_file = open('v3_conversion_log.txt', 'w', encoding='utf-8')
    
    def log(msg=""):
        print(msg)
        log_file.write(msg + "\n")
    
    log("=" * 70)
    log("KDP CONVERSION - VERSION 3")
    log("=" * 70)
    
    # Backup
    log(f"\nCreating backup...")
    shutil.copy2(SOURCE_FILE, BACKUP_FILE)
    log(f"  -> {BACKUP_FILE}")
    
    # Load
    log(f"\nLoading document...")
    doc = Document(SOURCE_FILE)
    log(f"  -> {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # 1. Remove TOC (do first before other processing)
    log(f"\n[1/7] Removing Table of Contents...")
    remove_toc(doc)
    log(f"  -> {stats['toc_removed']} TOC entries removed")
    
    # 2. Process tables (before emoji stripping to avoid issues)
    log(f"\n[2/7] Processing tables...")
    process_tables(doc)
    log(f"  -> {stats['tables_flattened']} standard tables flattened")
    log(f"  -> {stats['tables_to_bullets']} comparison tables to bullets")
    log(f"  -> {stats['tables_special']} special tables formatted")
    log(f"  -> {stats['checkmarks_replaced']} checkmarks replaced")
    
    # 3. Strip emoji from headings
    log(f"\n[3/7] Stripping emoji from headings...")
    strip_emoji_from_headings(doc)
    log(f"  -> {stats['headings_emoji_stripped']} heading runs cleaned")
    
    # 4. Strip emoji from body text
    log(f"\n[4/7] Stripping emoji from body text...")
    strip_emoji_from_body(doc)
    log(f"  -> {stats['body_emoji_stripped']} body runs cleaned")
    
    # 5. Fix callouts
    log(f"\n[5/7] Fixing callouts...")
    fix_callouts(doc)
    log(f"  -> {stats['callouts_fixed']} callouts standardized")
    
    # 6. Remove page breaks
    log(f"\n[6/7] Removing page breaks...")
    remove_page_breaks(doc)
    log(f"  -> {stats['page_breaks_removed']} page breaks removed")
    
    # 7. Center images & check alt text
    log(f"\n[7/7] Processing images...")
    center_images_and_check_alt(doc)
    log(f"  -> {stats['images_centered']} images centered")
    log(f"  -> {stats['images_missing_alt']} images were missing alt text (fixed)")
    
    # 8. Clear headers/footers
    log(f"\n[+] Clearing headers and footers...")
    clear_headers_footers(doc)
    log(f"  -> {stats['headers_cleared']} headers cleared")
    log(f"  -> {stats['footers_cleared']} footers cleared")
    
    # 9. Remove excess blanks
    log(f"\n[+] Removing excess blank paragraphs...")
    remove_excess_blanks(doc)
    log(f"  -> {stats['blank_paragraphs_removed']} blank paragraphs removed")
    log(f"  -> {stats['glossary_bullets_preserved']} glossary bullets preserved")
    
    # Save
    log(f"\nSaving KDP document...")
    doc.save(OUTPUT_FILE)
    log(f"  -> {OUTPUT_FILE}")
    
    # Summary
    log("\n" + "=" * 70)
    log("CONVERSION COMPLETE - SUMMARY")
    log("=" * 70)
    for key, value in stats.items():
        log(f"  {key.replace('_', ' ').title()}: {value}")
    
    log(f"\n  Output: {OUTPUT_FILE}")
    log(f"  Backup: {BACKUP_FILE}")
    
    log_file.close()


if __name__ == '__main__':
    main()
