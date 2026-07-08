from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from collections import Counter
import re

doc_path = r'public\reference\book\comptia-book.docx'

print("="*80)
print("DOCUMENT INSPECTION - COMPTIA-BOOK.DOCX")
print("="*80)

doc = Document(doc_path)

# ============================================================================
print("\n1. HEADING STRUCTURE & EMOJI ANALYSIS")
print("="*80)

headings_with_emoji = []
heading_styles = Counter()

for i, para in enumerate(doc.paragraphs):
    if para.style.name.startswith('Heading'):
        heading_styles[para.style.name] += 1
        
        # Check for emoji in text
        text = para.text
        # Common emoji patterns
        has_emoji = any(char for char in text if ord(char) > 0x1F000)
        
        if has_emoji:
            headings_with_emoji.append({
                'index': i,
                'style': para.style.name,
                'text': text[:80],
                'runs': len(para.runs)
            })

print(f"Heading Styles Found:")
for style, count in heading_styles.most_common():
    print(f"  {style}: {count}")

print(f"\nHeadings with Emoji: {len(headings_with_emoji)}")
print("\nFirst 10 examples:")
for item in headings_with_emoji[:10]:
    print(f"  [{item['index']}] {item['style']}: {item['text']}")
    print(f"    Runs: {item['runs']}")

# ============================================================================
print("\n2. CALLOUT PARAGRAPH ANALYSIS")
print("="*80)

callout_styles = [
    'Exam Trap',
    'Exam Tip', 
    'Pro Tip',
    'exam trap',
    'exam tip',
    'pro tip'
]

callouts_found = []
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name.lower() if para.style.name else ''
    
    if any(callout in style_name for callout in callout_styles):
        callouts_found.append({
            'index': i,
            'style': para.style.name,
            'text': para.text[:100],
            'starts_with_emoji': len(para.text) > 0 and ord(para.text[0]) > 0x1F000
        })

print(f"Callout Paragraphs Found: {len(callouts_found)}")
print("\nFirst 10 examples:")
for item in callouts_found[:10]:
    emoji_marker = "✓" if item['starts_with_emoji'] else "✗"
    print(f"  [{item['index']}] {emoji_marker} {item['style']}: {item['text'][:60]}")

# ============================================================================
print("\n3. TABLE ANALYSIS")
print("="*80)

print(f"Total Tables: {len(doc.tables)}")

table_analysis = []
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    
    # Check for merged cells
    has_merged = False
    merged_cells = []
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.tcPr
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    span_val = gridSpan.get(qn('w:val'))
                    has_merged = True
                    merged_cells.append(f"Row {row_idx}, Col {cell_idx}, Span={span_val}")
    
    # Get first row content for identification
    first_row_text = ' | '.join([cell.text[:30] for cell in table.rows[0].cells]) if rows > 0 else ''
    
    table_analysis.append({
        'index': i,
        'rows': rows,
        'cols': cols,
        'has_merged': has_merged,
        'merged_details': merged_cells[:3] if merged_cells else [],
        'first_row': first_row_text[:100]
    })

print(f"\nTable Summary:")
print(f"  2-column tables: {sum(1 for t in table_analysis if t['cols'] == 2)}")
print(f"  3-column tables: {sum(1 for t in table_analysis if t['cols'] == 3)}")
print(f"  4-column tables: {sum(1 for t in table_analysis if t['cols'] == 4)}")
print(f"  5+ column tables: {sum(1 for t in table_analysis if t['cols'] >= 5)}")
print(f"  Tables with merged cells: {sum(1 for t in table_analysis if t['has_merged'])}")

print("\nFirst 15 tables (with merge info):")
for t in table_analysis[:15]:
    merged_marker = "⚠️ MERGED" if t['has_merged'] else ""
    print(f"  Table {t['index']}: {t['rows']}x{t['cols']} {merged_marker}")
    print(f"    First row: {t['first_row'][:80]}")
    if t['merged_details']:
        print(f"    Merged: {', '.join(t['merged_details'])}")

print("\nTables with 5+ columns (likely image candidates):")
wide_tables = [t for t in table_analysis if t['cols'] >= 5]
for t in wide_tables[:10]:
    print(f"  Table {t['index']}: {t['rows']}x{t['cols']}")
    print(f"    {t['first_row'][:80]}")

# ============================================================================
print("\n4. IMAGE ANALYSIS")
print("="*80)

image_count = 0
images_without_alt = []

for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        if run._element.xpath('.//pic:pic'):
            image_count += 1
            
            # Check for alt text
            pics = run._element.xpath('.//pic:pic')
            for pic in pics:
                docPr = pic.xpath('.//wp:docPr')[0] if pic.xpath('.//wp:docPr') else None
                if docPr is not None:
                    alt_text = docPr.get('descr', '')
                    if not alt_text:
                        images_without_alt.append({
                            'para_index': i,
                            'para_text': para.text[:60] if para.text else '(empty)'
                        })

print(f"Total Images: {image_count}")
print(f"Images WITHOUT alt text: {len(images_without_alt)}")

if images_without_alt:
    print("\nFirst 10 images missing alt text:")
    for img in images_without_alt[:10]:
        print(f"  Paragraph {img['para_index']}: {img['para_text']}")

# ============================================================================
print("\n5. HEADER/FOOTER ANALYSIS")
print("="*80)

sections_with_headers = 0
sections_with_footers = 0

for section in doc.sections:
    if section.header.paragraphs:
        has_content = any(p.text.strip() for p in section.header.paragraphs)
        if has_content:
            sections_with_headers += 1
    
    if section.footer.paragraphs:
        has_content = any(p.text.strip() for p in section.footer.paragraphs)
        if has_content:
            sections_with_footers += 1

print(f"Total Sections: {len(doc.sections)}")
print(f"Sections with headers: {sections_with_headers}")
print(f"Sections with footers: {sections_with_footers}")

# ============================================================================
print("\n6. CONSECUTIVE EMPTY PARAGRAPHS")
print("="*80)

empty_runs = []
current_run = []

for i, para in enumerate(doc.paragraphs):
    if not para.text.strip():
        current_run.append(i)
    else:
        if len(current_run) >= 2:
            empty_runs.append((current_run[0], current_run[-1], len(current_run)))
        current_run = []

if current_run and len(current_run) >= 2:
    empty_runs.append((current_run[0], current_run[-1], len(current_run)))

print(f"Runs of 2+ consecutive empty paragraphs: {len(empty_runs)}")
if empty_runs:
    print("\nFirst 10 locations:")
    for start, end, count in empty_runs[:10]:
        print(f"  Paragraphs {start}-{end}: {count} consecutive empty")

# ============================================================================
print("\n7. GLOSSARY BULLET ANALYSIS")
print("="*80)

# Find glossary section
glossary_start = -1
for i, para in enumerate(doc.paragraphs):
    if 'glossary' in para.text.lower() and para.style.name.startswith('Heading'):
        glossary_start = i
        break

if glossary_start >= 0:
    print(f"Glossary found at paragraph {glossary_start}")
    
    # Check for colored bullets in glossary
    colored_bullets = []
    for i in range(glossary_start, min(glossary_start + 500, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        for run in para.runs:
            if '●' in run.text and run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                colored_bullets.append({
                    'para': i,
                    'text': para.text[:60],
                    'color': f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"
                })
                break
    
    print(f"Colored bullet entries found: {len(colored_bullets)}")
    if colored_bullets:
        print("\nFirst 10 colored bullets:")
        for item in colored_bullets[:10]:
            print(f"  Para {item['para']}: {item['color']} - {item['text']}")
else:
    print("Glossary section not found")

# ============================================================================
print("\n8. PARAGRAPH STYLE ANALYSIS")
print("="*80)

style_counts = Counter()
for para in doc.paragraphs:
    style_counts[para.style.name] += 1

print("Top 15 paragraph styles:")
for style, count in style_counts.most_common(15):
    print(f"  {style}: {count}")

# ============================================================================
print("\n" + "="*80)
print("INSPECTION COMPLETE")
print("="*80)

# Save detailed table analysis
with open('table_analysis.txt', 'w', encoding='utf-8') as f:
    f.write("DETAILED TABLE ANALYSIS\n")
    f.write("="*80 + "\n\n")
    
    for t in table_analysis:
        f.write(f"Table {t['index']}: {t['rows']} rows × {t['cols']} columns\n")
        f.write(f"  Has merged cells: {t['has_merged']}\n")
        if t['merged_details']:
            f.write(f"  Merge details: {', '.join(t['merged_details'])}\n")
        f.write(f"  First row: {t['first_row']}\n")
        f.write("\n")

print("\nDetailed table analysis saved to: table_analysis.txt")
