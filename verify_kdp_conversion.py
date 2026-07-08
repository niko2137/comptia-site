"""
Quick verification script to check KDP conversion results
Run after convert_to_kdp_complete.py finishes
"""

from docx import Document
import os

KDP_FILE = r'public\reference\book\comptia-book-KDP.docx'

def main():
    if not os.path.exists(KDP_FILE):
        print("ERROR: KDP file not found. Run convert_to_kdp_complete.py first.")
        return
    
    print("="*80)
    print("KDP CONVERSION VERIFICATION")
    print("="*80)
    
    doc = Document(KDP_FILE)
    
    # Check 1: Tables
    print(f"\n1. TABLES: {len(doc.tables)} remaining")
    if len(doc.tables) == 0:
        print("   OK All tables removed (flattened or converted to images)")
    else:
        print(f"   WARNING: {len(doc.tables)} tables still present")
        for i, table in enumerate(doc.tables[:5]):
            print(f"      Table {i}: {len(table.rows)}x{len(table.columns)}")
    
    # Check 2: Images
    image_count = 0
    centered_count = 0
    alt_text_count = 0
    
    for para in doc.paragraphs:
        pics = para._element.xpath('.//pic:pic')
        if pics:
            image_count += 1
            if para.alignment == 1:  # WD_PARAGRAPH_ALIGNMENT.CENTER
                centered_count += 1
            
            for pic in pics:
                docPr_list = pic.xpath('.//wp:docPr')
                if docPr_list and docPr_list[0].get('descr'):
                    alt_text_count += 1
    
    print(f"\n2. IMAGES: {image_count} total")
    print(f"   Centered: {centered_count}/{image_count}")
    print(f"   With alt text: {alt_text_count}/{image_count}")
    if image_count == centered_count == alt_text_count:
        print("   OK All images properly formatted")
    
    # Check 3: Headings with emoji
    headings_with_emoji = 0
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if any(ord(c) > 0x1F000 for c in para.text):
                headings_with_emoji += 1
                if headings_with_emoji <= 3:
                    print(f"      Found emoji in heading: {para.text[:60]}")
    
    print(f"\n3. HEADING EMOJI: {headings_with_emoji} headings still have emoji")
    if headings_with_emoji == 0:
        print("   OK All emoji removed from headings")
    else:
        print("   WARNING: Some headings still contain emoji")
    
    # Check 4: Headers/Footers
    headers_with_content = 0
    footers_with_content = 0
    
    for section in doc.sections:
        if any(p.text.strip() for p in section.header.paragraphs):
            headers_with_content += 1
        if any(p.text.strip() for p in section.footer.paragraphs):
            footers_with_content += 1
    
    print(f"\n4. HEADERS/FOOTERS:")
    print(f"   Headers with content: {headers_with_content}")
    print(f"   Footers with content: {footers_with_content}")
    if headers_with_content == 0 and footers_with_content == 0:
        print("   OK All headers and footers cleared")
    
    # Check 5: Excessive blanks
    consecutive_empty = []
    current_run = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            current_run.append(1)
        else:
            if len(current_run) >= 3:
                consecutive_empty.append(len(current_run))
            current_run = []
    
    if current_run and len(current_run) >= 3:
        consecutive_empty.append(len(current_run))
    
    print(f"\n5. BLANK PARAGRAPHS:")
    print(f"   Runs of 3+ consecutive: {len(consecutive_empty)}")
    if len(consecutive_empty) == 0:
        print("   OK No excessive blank paragraphs")
    else:
        print(f"   WARNING: Found {len(consecutive_empty)} sections with 3+ blanks")
    
    # Check 6: Glossary bullets
    glossary_bullets = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if '●' in run.text and run.font.color and run.font.color.rgb:
                glossary_bullets += 1
                break
    
    print(f"\n6. GLOSSARY BULLETS: {glossary_bullets} colored bullets found")
    if glossary_bullets > 0:
        print("   OK Glossary priority bullets preserved")
    
    # Summary
    print("\n" + "="*80)
    print("VERIFICATION SUMMARY")
    print("="*80)
    
    issues = 0
    
    if len(doc.tables) > 0:
        print(f"   WARNING: {len(doc.tables)} tables remain")
        issues += 1
    
    if headings_with_emoji > 0:
        print(f"   WARNING: {headings_with_emoji} headings still have emoji")
        issues += 1
    
    if headers_with_content > 0 or footers_with_content > 0:
        print(f"   WARNING: Headers/footers still have content")
        issues += 1
    
    if image_count != centered_count:
        print(f"   WARNING: Not all images centered")
        issues += 1
    
    if image_count != alt_text_count:
        print(f"   WARNING: Some images missing alt text")
        issues += 1
    
    if issues == 0:
        print("   OK No issues found - conversion successful!")
    else:
        print(f"   REVIEW: {issues} items need attention")
    
    print(f"\n   Document: {KDP_FILE}")
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print(f"   Images: {image_count}")
    print(f"   Sections: {len(doc.sections)}")

if __name__ == '__main__':
    main()
