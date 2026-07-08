"""
V3 INSPECTION SCRIPT
Inspects the master docx to understand:
- Table structure (headers, column counts, merged cells, embedded images)
- Styles used
- Image locations and alt text
- Headers/footers content
- Emoji usage in headings and body
- Callout patterns
- Glossary bullet structure
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import re
import json

SOURCE_FILE = r'public\reference\book\comptia-book.docx'
OUTPUT_FILE = 'inspection_results.txt'

# Emoji pattern (broad)
EMOJI_PATTERN = re.compile(
    '[\U0001F000-\U0001FFFF]|'
    '[\U00002600-\U000026FF]|'
    '[\U00002700-\U000027BF]|'
    '[\U0000FE00-\U0000FE0F]|'
    '[\U0000200D]|'
    '[\U00002300-\U000023FF]|'
    '[\U00002B00-\U00002BFF]|'
    '[\U0000203C-\U0000203C]|'
    '[\U00002049-\U00002049]|'
    '[\U000000A9]|'
    '[\U000000AE]|'
    '[\U0000231A-\U0000231B]|'
    '[\U000025AA-\U000025AB]|'
    '[\U000025B6]|'
    '[\U000025C0]|'
    '[\U000025FB-\U000025FE]'
)

def inspect_document():
    out = open(OUTPUT_FILE, 'w', encoding='utf-8')
    def p(text=""):
        out.write(text + "\n")
    
    p("=" * 80)
    p("V3 DOCUMENT INSPECTION")
    p("=" * 80)
    
    doc = Document(SOURCE_FILE)
    
    # Basic stats
    p(f"\nParagraphs: {len(doc.paragraphs)}")
    p(f"Tables: {len(doc.tables)}")
    p(f"Sections: {len(doc.sections)}")
    
    # === STYLES ===
    p("\n" + "=" * 80)
    p("STYLES IN USE")
    p("=" * 80)
    style_counts = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'None'
        style_counts[style_name] = style_counts.get(style_name, 0) + 1
    
    for style, count in sorted(style_counts.items(), key=lambda x: -x[1]):
        p(f"  {style}: {count}")
    
    # === TABLES ===
    p("\n" + "=" * 80)
    p("TABLE ANALYSIS")
    p("=" * 80)
    
    for i, table in enumerate(doc.tables):
        rows = table.rows
        cols = table.columns
        
        # Get header row text
        if len(rows) > 0:
            first_row_cells = rows[0].cells
            # Deduplicate merged cells
            seen = set()
            unique_headers = []
            for cell in first_row_cells:
                cell_id = id(cell._element)
                if cell_id not in seen:
                    seen.add(cell_id)
                    unique_headers.append(cell.text.strip()[:50])
            
            header_text = " | ".join(unique_headers)
        else:
            header_text = "(empty)"
        
        # Check for images in table
        table_xml = etree.tostring(table._element, encoding='unicode')
        has_images = 'pic:pic' in table_xml or 'w:drawing' in table_xml
        
        # Check for gridSpan (merged cells)
        has_merged = 'w:gridSpan' in table_xml
        
        # Check for checkmarks/x marks
        has_checkmarks = '\u2705' in table_xml or '\u274C' in table_xml or '\u2714' in table_xml
        
        # Get preceding paragraph (context for identification)
        table_elem = table._element
        prev_elem = table_elem.getprevious()
        preceding_text = ""
        if prev_elem is not None and prev_elem.tag == qn('w:p'):
            for t in prev_elem.itertext():
                preceding_text += t
            preceding_text = preceding_text.strip()[:80]
        
        num_rows = len(rows)
        num_unique_cols = len(unique_headers)
        
        # Get first 2 data rows for context
        sample_rows = []
        for row_idx in range(1, min(3, len(rows))):
            row_cells = rows[row_idx].cells
            seen_r = set()
            unique_r = []
            for cell in row_cells:
                cid = id(cell._element)
                if cid not in seen_r:
                    seen_r.add(cid)
                    unique_r.append(cell.text.strip()[:40])
            sample_rows.append(" | ".join(unique_r))
        
        p(f"\n  Table {i+1}/{len(doc.tables)}:")
        p(f"    Rows: {num_rows}, Unique Cols: {num_unique_cols}")
        p(f"    Headers: {header_text}")
        p(f"    Has Images: {has_images}")
        p(f"    Has Merged Cells: {has_merged}")
        p(f"    Has Checkmarks: {has_checkmarks}")
        p(f"    Preceding: {preceding_text}")
        for sr in sample_rows:
            p(f"    Sample Row: {sr}")
    
    # === IMAGES ===
    p("\n" + "=" * 80)
    p("IMAGE ANALYSIS")
    p("=" * 80)
    
    image_count = 0
    images_with_alt = 0
    images_without_alt = 0
    images_in_tables = 0
    missing_alt_contexts = []
    
    # Images in paragraphs
    for para in doc.paragraphs:
        drawings = para._element.findall('.//' + qn('w:drawing'))
        for drawing in drawings:
            image_count += 1
            # Check for alt text in docPr
            docPr = drawing.find('.//' + qn('wp:docPr'))
            if docPr is not None:
                alt = docPr.get('descr', '')
                if alt:
                    images_with_alt += 1
                else:
                    images_without_alt += 1
                    missing_alt_contexts.append(para.text[:60])
            else:
                images_without_alt += 1
                missing_alt_contexts.append(para.text[:60])
    
    # Images in tables
    for table in doc.tables:
        table_xml = etree.tostring(table._element, encoding='unicode')
        if 'w:drawing' in table_xml:
            drawings_in_table = table._element.findall('.//' + qn('w:drawing'))
            images_in_tables += len(drawings_in_table)
    
    p(f"\n  Total images in paragraphs: {image_count}")
    p(f"  Images with alt text: {images_with_alt}")
    p(f"  Images without alt text: {images_without_alt}")
    p(f"  Images embedded in tables: {images_in_tables}")
    
    if missing_alt_contexts:
        p(f"\n  Images missing alt text (context):")
        for ctx in missing_alt_contexts[:20]:
            p(f"    - near: '{ctx}'")
    
    # === EMOJI IN HEADINGS ===
    p("\n" + "=" * 80)
    p("EMOJI IN HEADINGS")
    p("=" * 80)
    
    emoji_headings = 0
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            text = para.text
            emojis_found = EMOJI_PATTERN.findall(text)
            if emojis_found:
                emoji_headings += 1
                if emoji_headings <= 30:
                    p(f"  [{para.style.name}] {text[:80]}")
                    p(f"    Emoji codepoints: {[hex(ord(e)) for e in emojis_found]}")
    
    p(f"\n  Total headings with emoji: {emoji_headings}")
    
    # === EMOJI IN BODY ===
    p("\n" + "=" * 80)
    p("EMOJI IN BODY TEXT (sample)")
    p("=" * 80)
    
    body_emoji_count = 0
    for para in doc.paragraphs:
        if para.style and not para.style.name.startswith('Heading'):
            text = para.text
            emojis_found = EMOJI_PATTERN.findall(text)
            if emojis_found:
                body_emoji_count += 1
                if body_emoji_count <= 15:
                    p(f"  [{para.style.name}] {text[:100]}")
                    p(f"    Emoji codepoints: {[hex(ord(e)) for e in emojis_found]}")
    
    p(f"\n  Total body paragraphs with emoji: {body_emoji_count}")
    
    # === CALLOUTS ===
    p("\n" + "=" * 80)
    p("CALLOUT STYLES")
    p("=" * 80)
    
    callout_styles = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if any(x in style_name.lower() for x in ['tip', 'trap', 'callout', 'note', 'warning', 'exam']):
            if style_name not in callout_styles:
                callout_styles[style_name] = []
            if len(callout_styles[style_name]) < 5:
                callout_styles[style_name].append(para.text[:120])
    
    for style, samples in callout_styles.items():
        p(f"\n  Style: '{style}' (showing {len(samples)} samples)")
        for s in samples:
            p(f"    -> {s}")
    
    # === HEADERS & FOOTERS ===
    p("\n" + "=" * 80)
    p("HEADERS & FOOTERS")
    p("=" * 80)
    
    for i, section in enumerate(doc.sections):
        header_text = ""
        for para in section.header.paragraphs:
            if para.text.strip():
                header_text += para.text.strip() + " | "
        
        footer_text = ""
        for para in section.footer.paragraphs:
            if para.text.strip():
                footer_text += para.text.strip() + " | "
        
        if header_text or footer_text:
            p(f"  Section {i+1}:")
            if header_text:
                p(f"    Header: {header_text[:100]}")
            if footer_text:
                p(f"    Footer: {footer_text[:100]}")
    
    if not any(True for s in doc.sections for pa in s.header.paragraphs if pa.text.strip()):
        p("  (No headers/footers with text found)")
    
    # === GLOSSARY BULLETS ===
    p("\n" + "=" * 80)
    p("GLOSSARY BULLETS (colored bullet)")
    p("=" * 80)
    
    bullet_colors = {}
    bullet_samples = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if '\u25cf' in run.text:  # ●
                try:
                    color = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'no-color'
                except:
                    color = 'error'
                bullet_colors[color] = bullet_colors.get(color, 0) + 1
                if color not in bullet_samples:
                    bullet_samples[color] = para.text[:60]
    
    for color, count in bullet_colors.items():
        p(f"  Color {color}: {count} bullets")
        if color in bullet_samples:
            p(f"    Sample: {bullet_samples[color]}")
    
    # === PAGE BREAKS ===
    p("\n" + "=" * 80)
    p("PAGE BREAKS")
    p("=" * 80)
    
    page_break_before_count = 0
    explicit_break_count = 0
    
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            if pPr.find(qn('w:pageBreakBefore')) is not None:
                page_break_before_count += 1
        
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    explicit_break_count += 1
    
    p(f"  pageBreakBefore: {page_break_before_count}")
    p(f"  Explicit page breaks (w:br type=page): {explicit_break_count}")
    
    # === TOC DETECTION ===
    p("\n" + "=" * 80)
    p("TOC DETECTION")
    p("=" * 80)
    
    toc_style_count = 0
    toc_heading_found = False
    for para in doc.paragraphs:
        if para.style and 'toc' in para.style.name.lower():
            toc_style_count += 1
            if toc_style_count <= 5:
                p(f"  TOC style paragraph: [{para.style.name}] {para.text[:80]}")
        if para.style and para.style.name.startswith('Heading'):
            if para.text.strip().lower() in ['table of contents', 'contents', 'toc']:
                toc_heading_found = True
                p(f"  TOC Heading found: '{para.text}'")
    
    p(f"\n  Total TOC-styled paragraphs: {toc_style_count}")
    p(f"  TOC heading found: {toc_heading_found}")
    
    p("\n" + "=" * 80)
    p("INSPECTION COMPLETE")
    p("=" * 80)
    
    out.close()
    print(f"Results written to {OUTPUT_FILE}")

if __name__ == '__main__':
    inspect_document()
