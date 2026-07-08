"""
V3 VERIFICATION SCRIPT
Re-opens the converted document and confirms:
- Zero tables remain
- All images are centered with alt text
- Headers/footers are empty
- No runs of 2+ blank paragraphs
- Glossary bullets retain colors
- No emoji in headings
- Callouts are standardized
- No page breaks remain
- Spot checks on flattened content
"""

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

OUTPUT_FILE = r'public\reference\book\comptia-book-KDP-v3.docx'
REPORT_FILE = 'v3_verification_report.txt'

# Emoji detection (same broad pattern)
EMOJI_PATTERN = re.compile(
    '[\U0001F000-\U0001FFFF]|'
    '[\U00002600-\U000026FF]|'
    '[\U00002700-\U000027BF]|'
    '[\U00002300-\U000023FF]|'
    '[\U00002B00-\U00002BFF]'
)

ALLOWED_EMOJI = {'\u2714', '\u2139', '\u26A0'}

def verify():
    out = open(REPORT_FILE, 'w', encoding='utf-8')
    def p(text=""):
        out.write(text + "\n")
    
    p("=" * 70)
    p("V3 VERIFICATION REPORT")
    p("=" * 70)
    
    doc = Document(OUTPUT_FILE)
    
    p(f"\nDocument loaded: {len(doc.paragraphs)} paragraphs")
    
    issues = []
    
    # === CHECK 1: Zero tables ===
    p("\n--- CHECK 1: Tables ---")
    table_count = len(doc.tables)
    if table_count == 0:
        p("  PASS: Zero tables remain")
    else:
        p(f"  FAIL: {table_count} tables still exist!")
        issues.append(f"{table_count} tables remain")
    
    # === CHECK 2: Images centered with alt text ===
    p("\n--- CHECK 2: Images ---")
    images_total = 0
    images_centered = 0
    images_with_alt = 0
    images_missing_alt = []
    
    for para in doc.paragraphs:
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            images_total += 1
            
            # Check centering
            if para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                images_centered += 1
            
            # Check alt text
            for drawing in drawings:
                docPr = drawing.find('.//' + qn('wp:docPr'))
                if docPr is not None:
                    alt = docPr.get('descr', '')
                    if alt:
                        images_with_alt += 1
                    else:
                        images_missing_alt.append(para.text[:50])
                else:
                    images_missing_alt.append(para.text[:50])
    
    p(f"  Total images: {images_total}")
    p(f"  Centered: {images_centered}")
    p(f"  With alt text: {images_with_alt}")
    if images_missing_alt:
        p(f"  FAIL: {len(images_missing_alt)} images missing alt text")
        issues.append(f"{len(images_missing_alt)} images missing alt")
    else:
        p("  PASS: All images have alt text")
    if images_centered == images_total:
        p("  PASS: All images centered")
    else:
        p(f"  WARN: {images_total - images_centered} images not centered")
    
    # === CHECK 3: Headers/Footers empty ===
    p("\n--- CHECK 3: Headers & Footers ---")
    headers_with_text = 0
    footers_with_text = 0
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                headers_with_text += 1
        for para in section.footer.paragraphs:
            if para.text.strip():
                footers_with_text += 1
    
    if headers_with_text == 0 and footers_with_text == 0:
        p("  PASS: All headers/footers are empty")
    else:
        p(f"  FAIL: {headers_with_text} headers, {footers_with_text} footers still have text")
        issues.append("Headers/footers not cleared")
    
    # === CHECK 4: No runs of 2+ blank paragraphs ===
    p("\n--- CHECK 4: Blank Paragraph Runs ---")
    consecutive = 0
    max_consecutive = 0
    violations = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            consecutive += 1
            if consecutive > max_consecutive:
                max_consecutive = consecutive
            if consecutive >= 2:
                violations += 1
        else:
            consecutive = 0
    
    if max_consecutive <= 1:
        p("  PASS: No runs of 2+ blank paragraphs")
    else:
        p(f"  FAIL: Max consecutive blanks = {max_consecutive}, violations = {violations}")
        issues.append(f"Blank paragraph runs (max {max_consecutive})")
    
    # === CHECK 5: Glossary bullets ===
    p("\n--- CHECK 5: Glossary Bullets ---")
    bullet_colors = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if '\u25cf' in run.text:
                try:
                    color = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'no-color'
                except:
                    color = 'error'
                bullet_colors[color] = bullet_colors.get(color, 0) + 1
    
    if bullet_colors:
        p(f"  Glossary bullets found:")
        for color, count in bullet_colors.items():
            p(f"    {color}: {count}")
        if 'FF0000' in bullet_colors and 'FFC000' in bullet_colors and '00B300' in bullet_colors:
            p("  PASS: All three priority colors preserved")
        else:
            p("  WARN: Some colors may be missing")
            issues.append("Glossary bullet colors incomplete")
    else:
        p("  FAIL: No glossary bullets found!")
        issues.append("Glossary bullets lost")
    
    # === CHECK 6: No emoji in headings ===
    p("\n--- CHECK 6: Emoji in Headings ---")
    headings_with_emoji = 0
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            text = para.text
            for char in text:
                cp = ord(char)
                if (0x1F000 <= cp <= 0x1FFFF or
                    0x2600 <= cp <= 0x26FF or
                    0x2700 <= cp <= 0x27BF or
                    0x2300 <= cp <= 0x23FF or
                    0x2B00 <= cp <= 0x2BFF):
                    if char not in ALLOWED_EMOJI:
                        headings_with_emoji += 1
                        if headings_with_emoji <= 5:
                            p(f"    Still has emoji: [{para.style.name}] {text[:60]}")
                        break
    
    if headings_with_emoji == 0:
        p("  PASS: No emoji in headings")
    else:
        p(f"  FAIL: {headings_with_emoji} headings still have emoji")
        issues.append(f"{headings_with_emoji} headings with emoji")
    
    # === CHECK 7: Callouts ===
    p("\n--- CHECK 7: Callouts ---")
    callout_issues = 0
    for para in doc.paragraphs:
        if not para.style:
            continue
        style = para.style.name
        text = para.text
        
        if style == 'Tip -Trap':
            if not text.startswith('\u26A0'):
                callout_issues += 1
                if callout_issues <= 3:
                    p(f"    Bad Trap callout: {text[:80]}")
        elif style == 'Tip -Exam':
            if not text.startswith('\u2139'):
                callout_issues += 1
                if callout_issues <= 3:
                    p(f"    Bad Exam callout: {text[:80]}")
        elif style == 'Tip -Pro':
            if not text.startswith('\u2714'):
                callout_issues += 1
                if callout_issues <= 3:
                    p(f"    Bad Pro callout: {text[:80]}")
    
    if callout_issues == 0:
        p("  PASS: All callouts standardized")
    else:
        p(f"  FAIL: {callout_issues} callouts not standardized")
        issues.append(f"{callout_issues} callout issues")
    
    # === CHECK 8: Page breaks ===
    p("\n--- CHECK 8: Page Breaks ---")
    page_breaks = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            if pPr.find(qn('w:pageBreakBefore')) is not None:
                page_breaks += 1
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    page_breaks += 1
    
    if page_breaks == 0:
        p("  PASS: No page breaks remain")
    else:
        p(f"  FAIL: {page_breaks} page breaks remain")
        issues.append(f"{page_breaks} page breaks remain")
    
    # === CHECK 9: Spot checks ===
    p("\n--- CHECK 9: Spot Checks (content samples) ---")
    
    # Look for motherboard content
    motherboard_found = False
    for para in doc.paragraphs:
        if 'CPU Socket LGA' in para.text or 'Board Components' in para.text:
            motherboard_found = True
            p(f"  Motherboard content found: {para.text[:80]}")
            break
    
    if motherboard_found:
        p("  PASS: Motherboard Components table content preserved")
    else:
        p("  WARN: Could not find Motherboard Components content")
    
    # Look for bullet-formatted comparison content
    bullet_found = False
    for para in doc.paragraphs:
        if para.text.startswith('\u2022') and ('TCP' in para.text or 'UDP' in para.text):
            bullet_found = True
            p(f"  Bullet comparison found: {para.text[:80]}")
            break
    
    if bullet_found:
        p("  PASS: Comparison table bullet format present")
    else:
        p("  WARN: Could not find bullet comparison content")
    
    # Look for IETAID content
    ietaid_found = False
    for para in doc.paragraphs:
        if 'Step 1 - I' in para.text or 'IETAID' in para.text:
            ietaid_found = True
            p(f"  IETAID content found: {para.text[:80]}")
            break
    
    if ietaid_found:
        p("  PASS: IETAID table content preserved")
    else:
        p("  WARN: Could not find IETAID content")
    
    # Look for Windows Edition YES/NO replacement
    yes_no_found = False
    for para in doc.paragraphs:
        if 'BitLocker' in para.text and ('YES' in para.text or 'NO' in para.text):
            yes_no_found = True
            p(f"  Windows Edition YES/NO: {para.text[:80]}")
            break
    
    if yes_no_found:
        p("  PASS: Windows Edition checkmarks replaced with YES/NO")
    else:
        p("  WARN: Could not find Windows Edition YES/NO content")
    
    # === SUMMARY ===
    p("\n" + "=" * 70)
    p("VERIFICATION SUMMARY")
    p("=" * 70)
    
    if not issues:
        p("\n  ALL CHECKS PASSED!")
    else:
        p(f"\n  {len(issues)} ISSUE(S) FOUND:")
        for issue in issues:
            p(f"    - {issue}")
    
    p(f"\n  Total paragraphs in output: {len(doc.paragraphs)}")
    p(f"  Total images preserved: {images_total}")
    
    out.close()
    print(f"Verification report written to {REPORT_FILE}")


if __name__ == '__main__':
    verify()
