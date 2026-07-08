"""
COMPREHENSIVE AUDIT SCRIPT
Audits both KDP-final.docx and Interior-paper_and_hardback-FINAL-backup.docx
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree
import re
from collections import Counter, defaultdict

FILES = {
    'KDP': r'public\reference\book\KDP-final.docx',
    'HARDBACK': r'public\reference\book\Interior-paper_and_hardback-FINAL-backup.docx',
}
REPORT = 'COMPLETE_AUDIT_REPORT.txt'
out = None

def p(text=""):
    out.write(text + "\n")

ALLOWED_EMOJI = {'\u2714', '\u2139', '\u26A0'}

def is_emoji(char):
    cp = ord(char)
    if char in ALLOWED_EMOJI: return False
    if 0x1F000 <= cp <= 0x1FFFF: return True
    if 0x2600 <= cp <= 0x26FF and char not in ALLOWED_EMOJI: return True
    if 0x2700 <= cp <= 0x27BF and char != '\u2714': return True
    if 0x2300 <= cp <= 0x23FF: return True
    if 0x2B00 <= cp <= 0x2BFF: return True
    if 0xFE00 <= cp <= 0xFE0F: return True
    if cp == 0x200D: return True
    return False

# Technical terms that spell-checkers would flag but are valid
TECH_TERMS = {
    'comptia', 'wifi', 'bluetooth', 'usb', 'hdmi', 'sata', 'nvme',
    'pcie', 'ddr', 'dimm', 'sodimm', 'uefi', 'bios', 'apipa',
    'dhcp', 'dns', 'tcp', 'udp', 'smtp', 'imap', 'pop3',
    'ssid', 'wpa', 'wep', 'aes', 'tkip', 'vpn', 'vlan',
    'nat', 'pat', 'ldap', 'rdp', 'vnc', 'ssh', 'sftp',
    'ftp', 'http', 'https', 'ssl', 'tls', 'ipsec', 'ipv4',
    'ipv6', 'subnet', 'subnetting', 'cidr', 'mbps', 'gbps',
    'ghz', 'mhz', 'tb', 'gb', 'mb', 'ntfs', 'fat32', 'exfat',
    'gpt', 'mbr', 'raid', 'ssd', 'hdd', 'nand', 'sshd',
    'gpu', 'cpu', 'ram', 'rom', 'psu', 'atx', 'matx', 'itx',
    'lga', 'pga', 'bga', 'amd', 'intel', 'ryzen', 'iot',
    'zigbee', 'zwave', 'mqtt', 'slaac', 'eui', 'imei', 'imsi',
    'iccid', 'sim', 'esim', 'nfc', 'gps', 'lte', 'wwan',
    'macos', 'linux', 'ios', 'android', 'powershell', 'cmd',
    'winre', 'bitlocker', 'hyper', 'hypervisor', 'vmware',
    'virtualbox', 'saas', 'paas', 'iaas', 'mdm', 'byod',
    'soho', 'osi', 'dora', 'ietaid', 'pbq', 'pbqs',
    'displayport', 'thunderbolt', 'usbc', 'ethernet',
    'rj45', 'utp', 'stp', 'sfp', 'qsfp', 'poe',
    'oled', 'lcd', 'led', 'ips', 'tn', 'va',
    'inkjet', 'laserjet', 'multifunction', 'mfp',
    'tpm', 'esd', 'msds', 'sds', 'hvac', 'ups',
    'kvm', 'ip', 'mac', 'arp', 'icmp', 'igmp',
    'tracert', 'traceroute', 'ipconfig', 'ifconfig',
    'nslookup', 'netstat', 'pathping', 'gpupdate',
    'sfc', 'dism', 'chkdsk', 'diskpart', 'robocopy',
    'xcopy', 'chmod', 'chown', 'grep', 'sudo', 'apt',
    'yum', 'nano', 'vim', 'dd', 'df', 'du', 'ps',
    'cron', 'systemctl', 'journalctl', 'dmesg',
    'ext4', 'xfs', 'apfs', 'hfs', 'ntfs',
    'msc', 'mmc', 'gpedit', 'regedit', 'msconfig',
    'devmgmt', 'diskmgmt', 'compmgmt', 'perfmon',
    'resmon', 'eventvwr', 'lusrmgr', 'secpol',
    'onedrive', 'icloud', 'carddav', 'caldav',
    'activesync', 'oauth', 'mfa', 'totp',
    'pii', 'phi', 'gdpr', 'hipaa', 'pci', 'dss',
    'itam', 'rfc', 'sla', 'mttr', 'mttf', 'mtbf',
    'ai', 'ml', 'llm', 'chatgpt', 'genai', 'chatbot',
    'ransomware', 'malware', 'rootkit', 'keylogger',
    'trojan', 'botnet', 'phishing', 'spearphishing',
    'vishing', 'smishing', 'whaling', 'tailgating',
    'dumpster', 'baiting', 'pretexting', 'qrishing',
    'wps', 'wpa2', 'wpa3', 'ccmp', 'sae',
}

def detect_merged_words(text):
    """Detect words that appear merged (camelCase-like in non-code context)."""
    # Pattern: lowercase followed by uppercase in middle of word
    merged = re.findall(r'[a-z][A-Z][a-z]', text)
    # Filter out legitimate camelCase terms
    legit = {'MHz', 'GHz', 'mATX', 'eReader', 'iCloud', 'macOS', 'iOS',
             'iPadOS', 'watchOS', 'tvOS', 'iPhone', 'iPad', 'iPod',
             'BitLocker', 'PowerShell', 'VirtualBox', 'TeamViewer',
             'AnyDesk', 'JavaScript', 'GitHub', 'YouTube', 'WiFi',
             'FireOS', 'ChromeOS', 'exFAT', 'CardDAV', 'CalDAV',
             'ActiveSync', 'FileVault', 'AirDrop', 'TimeMachine',
             'EverNote', 'OneNote', 'OneDrive', 'InPrivate',
             'QuickBooks', 'AutoRun', 'AutoPlay', 'McAfee',
             'WinRE', 'BitLocker', 'ReadyBoost', 'EasyBCD',
             'PageFile', 'HyperV'}
    
    results = []
    # Find patterns like "wordWord" that aren't in legit list
    pattern = re.compile(r'\b(\w*[a-z][A-Z]\w*)\b')
    for match in pattern.finditer(text):
        word = match.group()
        if word not in legit and not any(t in word for t in legit):
            # Check if it's a known merged issue like "ToUnderstand"
            results.append(word)
    return results


def audit_document(doc, name):
    """Run all audit checks on a single document."""
    p(f"\n{'='*70}")
    p(f"AUDITING: {name}")
    p(f"{'='*70}")
    p(f"Paragraphs: {len(doc.paragraphs)}")
    p(f"Tables: {len(doc.tables)}")
    p(f"Sections: {len(doc.sections)}")
    
    issues = []
    warnings = []
    
    # ============================================
    # PART 1: STRUCTURE & FORMATTING
    # ============================================
    p(f"\n{'─'*50}")
    p("PART 1: STRUCTURE, FORMATTING, SPACING")
    p(f"{'─'*50}")
    
    # Tables check
    if doc.tables:
        issues.append(f"{len(doc.tables)} tables remain (should be 0 for KDP)")
        p(f"  ❌ {len(doc.tables)} tables found")
    else:
        p(f"  ✓ Zero tables")
    
    # Page breaks
    page_breaks = 0
    for para in doc.paragraphs:
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            if pPr.find(qn('w:pageBreakBefore')) is not None:
                page_breaks += 1
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    page_breaks += 1
    
    if name == 'KDP' and page_breaks > 0:
        issues.append(f"{page_breaks} page breaks (should be 0 for KDP)")
        p(f"  ❌ {page_breaks} page breaks")
    elif name == 'HARDBACK':
        p(f"  ℹ {page_breaks} page breaks (expected for print)")
    else:
        p(f"  ✓ No page breaks")
    
    # Headers/footers
    hf_count = 0
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                hf_count += 1
        for para in section.footer.paragraphs:
            if para.text.strip():
                hf_count += 1
    
    if name == 'KDP' and hf_count > 0:
        issues.append(f"{hf_count} headers/footers with content (remove for KDP)")
        p(f"  ❌ {hf_count} headers/footers with text")
    elif name == 'HARDBACK':
        p(f"  ℹ {hf_count} headers/footers (expected for print)")
    else:
        p(f"  ✓ Headers/footers clear")
    
    # Double spaces
    double_space_paras = []
    for i, para in enumerate(doc.paragraphs):
        if '  ' in para.text:
            double_space_paras.append((i, para.text[:60]))
    
    if double_space_paras:
        warnings.append(f"{len(double_space_paras)} paragraphs with double spaces")
        p(f"  ⚠ {len(double_space_paras)} paragraphs with double spaces")
        for idx, text in double_space_paras[:5]:
            p(f"      Para {idx}: '{text}'")
    else:
        p(f"  ✓ No double spaces")
    
    # Leading whitespace
    leading_ws = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if text and (text[0] == ' ' or text[0] == '\t'):
            leading_ws.append((i, para.style.name if para.style else 'None', text[:50]))
    
    if leading_ws:
        warnings.append(f"{len(leading_ws)} paragraphs with leading whitespace")
        p(f"  ⚠ {len(leading_ws)} paragraphs with leading whitespace")
        for idx, style, text in leading_ws[:5]:
            p(f"      Para {idx} [{style}]: '{text}'")
    else:
        p(f"  ✓ No leading whitespace")
    
    # Consecutive blank paragraphs
    max_blanks = 0
    current_blanks = 0
    blank_locations = []
    for i, para in enumerate(doc.paragraphs):
        has_img = bool(para._element.findall('.//' + qn('w:drawing')))
        if not para.text.strip() and not has_img:
            current_blanks += 1
            if current_blanks > max_blanks:
                max_blanks = current_blanks
            if current_blanks == 3:
                blank_locations.append(i)
        else:
            current_blanks = 0
    
    if max_blanks >= 3:
        warnings.append(f"Max {max_blanks} consecutive blank paragraphs")
        p(f"  ⚠ Max {max_blanks} consecutive blanks (at paras: {blank_locations[:5]})")
    else:
        p(f"  ✓ No excessive blank runs (max {max_blanks})")
    
    # Hidden/invisible characters
    hidden_chars = defaultdict(int)
    for para in doc.paragraphs:
        for run in para.runs:
            for char in run.text:
                cp = ord(char)
                if cp == 0x200B: hidden_chars['ZWSP'] += 1
                elif cp == 0x00AD: hidden_chars['SoftHyphen'] += 1
                elif cp == 0xFEFF: hidden_chars['BOM'] += 1
                elif cp == 0x2028: hidden_chars['LineSep'] += 1
                elif cp == 0x2029: hidden_chars['ParaSep'] += 1
    
    if hidden_chars:
        warnings.append(f"Hidden characters: {dict(hidden_chars)}")
        p(f"  ⚠ Hidden characters found: {dict(hidden_chars)}")
    else:
        p(f"  ✓ No hidden/invisible characters")
    
    return issues, warnings


def audit_spelling_and_merging(doc, name):
    """Part 2: Detect merged words and potential misspellings."""
    p(f"\n{'─'*50}")
    p("PART 2: WORD MERGING & SPELLING ISSUES")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    
    # Merged words detection
    merged_found = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        merged = detect_merged_words(text)
        for word in merged:
            # Filter out common false positives
            if len(word) > 3 and word.lower() not in TECH_TERMS:
                merged_found.append((i, word, text[:60]))
    
    if merged_found:
        # Deduplicate by word
        unique_merged = {}
        for idx, word, context in merged_found:
            if word not in unique_merged:
                unique_merged[word] = (idx, context)
        
        issues.append(f"{len(unique_merged)} unique merged words found")
        p(f"  ❌ {len(unique_merged)} merged words detected:")
        for word, (idx, context) in list(unique_merged.items())[:20]:
            p(f"      '{word}' (para {idx}): {context}")
    else:
        p(f"  ✓ No merged words detected")
    
    # Common misspelling patterns
    misspell_patterns = [
        (r'\brecieve\b', 'receive'),
        (r'\boccured\b', 'occurred'),
        (r'\bseperete\b', 'separate'),
        (r'\bneccessary\b', 'necessary'),
        (r'\benviroment\b', 'environment'),
        (r'\bperformace\b', 'performance'),
        (r'\bconfiguartion\b', 'configuration'),
        (r'\bauthentiction\b', 'authentication'),
        (r'\btroubleshooing\b', 'troubleshooting'),
        (r'\bencryptoin\b', 'encryption'),
        (r'\bconnecton\b', 'connection'),
        (r'\bnetowrk\b', 'network'),
        (r'\boperting\b', 'operating'),
        (r'\bsytems\b', 'systems'),
        (r'\bmanagment\b', 'management'),
        (r'\bvirtualiztion\b', 'virtualization'),
        (r'\btechnolgy\b', 'technology'),
        (r'\binfromation\b', 'information'),
        (r'\bsercurity\b', 'security'),
        (r'\bdestionation\b', 'destination'),
        (r'\bcompatabile\b', 'compatible'),
        (r'\baccesible\b', 'accessible'),
        (r'\bthier\b', 'their'),
        (r'\bteh\b', 'the'),
        (r'\bwich\b', 'which'),
        (r'\bbeacuse\b', 'because'),
        (r'\bwindwos\b', 'Windows'),
        (r'\bmotherbaord\b', 'motherboard'),
        (r'\bprocesssor\b', 'processor'),
    ]
    
    spelling_issues = []
    for para in doc.paragraphs:
        text = para.text
        for pattern, correct in misspell_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                spelling_issues.append((matches[0], correct, text[:60]))
    
    if spelling_issues:
        issues.append(f"{len(spelling_issues)} potential misspellings")
        p(f"  ❌ {len(spelling_issues)} potential misspellings:")
        for wrong, correct, ctx in spelling_issues[:10]:
            p(f"      '{wrong}' → '{correct}': {ctx}")
    else:
        p(f"  ✓ No common misspellings detected")
    
    # Check for missing spaces after periods
    missing_space_after_period = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        # Pattern: period followed by uppercase without space (not abbreviations)
        matches = re.findall(r'[a-z]\.[A-Z]', text)
        if matches:
            # Filter out abbreviations like "e.g." or "i.e."
            real_issues = [m for m in matches if m[0] not in 'eg']
            if real_issues:
                missing_space_after_period.append((i, text[:80]))
    
    if missing_space_after_period:
        issues.append(f"{len(missing_space_after_period)} missing spaces after periods")
        p(f"  ❌ {len(missing_space_after_period)} missing space after period:")
        for idx, text in missing_space_after_period[:10]:
            p(f"      Para {idx}: {text}")
    else:
        p(f"  ✓ No missing spaces after periods")
    
    return issues, warnings


def audit_images(doc, name):
    """Part 3: Image audit - centering, alt text, names."""
    p(f"\n{'─'*50}")
    p("PART 3: IMAGE AUDIT")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    
    img_details = []
    for i, para in enumerate(doc.paragraphs):
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            is_centered = para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            for d in drawings:
                docPr = d.find('.//' + qn('wp:docPr'))
                img_name = ''
                alt_text = ''
                if docPr is not None:
                    img_name = docPr.get('name', '')
                    alt_text = docPr.get('descr', '')
                
                # Get context
                prev_text = doc.paragraphs[i-1].text[:50] if i > 0 else ''
                next_text = doc.paragraphs[i+1].text[:50] if i < len(doc.paragraphs)-1 else ''
                
                img_details.append({
                    'para_idx': i,
                    'name': img_name,
                    'alt': alt_text,
                    'centered': is_centered,
                    'prev': prev_text,
                    'next': next_text,
                    'style': para.style.name if para.style else 'None'
                })
    
    p(f"  Total images found: {len(img_details)}")
    
    # Check centering
    not_centered = [img for img in img_details if not img['centered']]
    if not_centered:
        issues.append(f"{len(not_centered)} images not centered")
        p(f"  ❌ {len(not_centered)} images NOT centered:")
        for img in not_centered[:5]:
            p(f"      Para {img['para_idx']} '{img['name']}' near: {img['prev'][:40]}")
    else:
        p(f"  ✓ All images centered")
    
    # Check alt text
    no_alt = [img for img in img_details if not img['alt']]
    if no_alt:
        issues.append(f"{len(no_alt)} images without alt text")
        p(f"  ❌ {len(no_alt)} images missing alt text:")
        for img in no_alt[:5]:
            p(f"      Para {img['para_idx']} '{img['name']}' near: {img['prev'][:40]}")
    else:
        p(f"  ✓ All images have alt text")
    
    # List all images with details
    p(f"\n  Image inventory:")
    for img in img_details:
        centered_mark = '✓' if img['centered'] else '✗'
        alt_mark = '✓' if img['alt'] else '✗'
        p(f"    [{centered_mark}C {alt_mark}A] {img['name'][:40]:40s} | Alt: {img['alt'][:40]}")
    
    return issues, warnings


def audit_content_accuracy(doc, name):
    """Part 4: CompTIA A+ 220-1201/1202 content accuracy."""
    p(f"\n{'─'*50}")
    p("PART 4: COMPTIA A+ CONTENT ACCURACY (220-1201/1202)")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    all_text = '\n'.join(para.text for para in doc.paragraphs)
    
    # Key facts that must be correct for CompTIA A+ v15
    fact_checks = [
        # Exam details
        ('220-1201', 'Core 1 exam number should be 220-1201'),
        ('220-1202', 'Core 2 exam number should be 220-1202'),
        ('90 minutes', 'Each exam is 90 minutes'),
        ('90 questions', 'Each exam is max 90 questions'),
        ('675', 'Passing score is 675 for Core 1'),
        ('700', 'Passing score is 700 for Core 2'),
        
        # Port numbers (critical)
        ('port 20', 'FTP data port'),
        ('port 21', 'FTP control port'),
        ('port 22', 'SSH/SFTP port'),
        ('port 23', 'Telnet port'),
        ('port 25', 'SMTP port'),
        ('port 53', 'DNS port'),
        ('port 80', 'HTTP port'),
        ('port 110', 'POP3 port'),
        ('port 143', 'IMAP port'),
        ('port 443', 'HTTPS port'),
        ('port 445', 'SMB port'),
        ('port 3389', 'RDP port'),
        
        # Key protocols
        ('WPA3', 'Latest Wi-Fi security standard'),
        ('Wi-Fi 6E', 'Wi-Fi 6E (6 GHz band)'),
        ('802.11ax', '802.11ax standard'),
        ('USB4', 'USB4 standard'),
        ('Thunderbolt', 'Thunderbolt standard'),
        ('DDR5', 'DDR5 RAM'),
        ('PCIe 4.0', 'PCIe 4.0'),
        ('NVMe', 'NVMe storage'),
        ('TPM 2.0', 'TPM 2.0 requirement'),
        
        # OS details
        ('Windows 11', 'Windows 11 covered'),
        ('Windows 10', 'Windows 10 covered'),
        
        # Key processes
        ('DORA', 'DHCP DORA process'),
        ('7-step', '7-step malware removal OR 7 layers'),
        ('6-step', '6-step troubleshooting'),
    ]
    
    p("  Checking critical facts present in document:")
    missing_facts = []
    for search, description in fact_checks:
        if search.lower() not in all_text.lower():
            missing_facts.append(description)
            p(f"    ❌ NOT FOUND: {description} (searched: '{search}')")
        else:
            count = all_text.lower().count(search.lower())
            p(f"    ✓ {description} (found {count}x)")
    
    if missing_facts:
        issues.append(f"{len(missing_facts)} critical facts missing")
    
    # Check for outdated information
    p("\n  Checking for outdated/incorrect info:")
    outdated_checks = [
        ('220-1001', 'Old exam number 220-1001 (should be 220-1201)'),
        ('220-1002', 'Old exam number 220-1002 (should be 220-1202)'),
        ('220-1101', 'Previous exam number 220-1101 (should be 220-1201)'),
        ('220-1102', 'Previous exam number 220-1102 (should be 220-1202)'),
        ('Windows 8', 'Windows 8 (no longer on exam)'),
        ('Windows 7', 'Windows 7 (no longer on exam)'),
    ]
    
    for search, description in outdated_checks:
        if search.lower() in all_text.lower():
            # Check context - might be mentioned historically
            count = all_text.lower().count(search.lower())
            warnings.append(f"'{search}' found {count}x - verify context")
            p(f"    ⚠ FOUND '{search}' ({count}x) - {description}")
        else:
            p(f"    ✓ No '{search}' found")
    
    # Domain weight verification
    p("\n  Checking exam domain weights:")
    domain_weights_core1 = {
        'Mobile Devices': '13%',
        'Networking': '23%',
        'Hardware': '27%',
        'Virtualization': '12%',
        'Troubleshooting': '25%',
    }
    domain_weights_core2 = {
        'Operating Systems': '28%',
        'Security': '28%',
        'Software Troubleshooting': '16%',
        'Operational Procedures': '28%',
    }
    
    for domain, weight in domain_weights_core1.items():
        if weight in all_text:
            p(f"    ✓ Core 1 - {domain}: {weight}")
        else:
            warnings.append(f"Core 1 {domain} weight {weight} not found explicitly")
            p(f"    ⚠ Core 1 - {domain}: {weight} not found as-is (verify)")
    
    for domain, weight in domain_weights_core2.items():
        if weight in all_text:
            p(f"    ✓ Core 2 - {domain}: {weight}")
        else:
            warnings.append(f"Core 2 {domain} weight {weight} not found")
            p(f"    ⚠ Core 2 - {domain}: {weight} not found as-is (verify)")
    
    return issues, warnings


def audit_questions_answers(doc, name):
    """Part 5: Practice questions and answers alignment."""
    p(f"\n{'─'*50}")
    p("PART 5: PRACTICE QUESTIONS & ANSWERS")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    
    # Find question patterns
    questions = []
    answers = []
    in_questions = False
    in_answers = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Detect Q&A sections
        if 'Practice Test' in text and para.style and 'Heading' in para.style.name:
            in_questions = True
            in_answers = False
            continue
        if 'Answer' in text and 'Key' in text and para.style and 'Heading' in para.style.name:
            in_questions = False
            in_answers = True
            continue
        if para.style and 'Heading 1' in para.style.name and in_answers:
            in_answers = False
            continue
        
        # Collect questions (numbered)
        if in_questions:
            q_match = re.match(r'^(\d+)\.\s+(.+)', text)
            if q_match:
                questions.append((int(q_match.group(1)), q_match.group(2)[:60], i))
        
        # Collect answers
        if in_answers:
            a_match = re.match(r'^(\d+)\.\s+([A-D])', text)
            if a_match:
                answers.append((int(a_match.group(1)), a_match.group(2), i))
    
    p(f"  Questions found: {len(questions)}")
    p(f"  Answers found: {len(answers)}")
    
    if questions and answers:
        # Check alignment
        q_numbers = [q[0] for q in questions]
        a_numbers = [a[0] for a in answers]
        
        # Check for gaps in question numbering
        if q_numbers:
            expected = list(range(1, max(q_numbers) + 1))
            missing_qs = set(expected) - set(q_numbers)
            if missing_qs:
                issues.append(f"Missing question numbers: {sorted(missing_qs)[:10]}")
                p(f"  ❌ Missing questions: {sorted(missing_qs)[:10]}")
        
        # Check Q/A count match
        if len(questions) != len(answers):
            issues.append(f"Q/A mismatch: {len(questions)} questions vs {len(answers)} answers")
            p(f"  ❌ Mismatch: {len(questions)} questions vs {len(answers)} answers")
        else:
            p(f"  ✓ Q/A counts match ({len(questions)} each)")
        
        # Check answer values are valid (A-D)
        invalid_answers = [a for a in answers if a[1] not in 'ABCD']
        if invalid_answers:
            issues.append(f"{len(invalid_answers)} invalid answer choices")
            p(f"  ❌ Invalid answers: {invalid_answers[:5]}")
        else:
            p(f"  ✓ All answers are valid (A-D)")
        
        # Check for duplicate question numbers
        q_counter = Counter(q_numbers)
        dupes = {num: count for num, count in q_counter.items() if count > 1}
        if dupes:
            warnings.append(f"Duplicate question numbers: {dupes}")
            p(f"  ⚠ Duplicate Q numbers: {dupes}")
    else:
        if not questions:
            warnings.append("No numbered questions found")
            p(f"  ⚠ No numbered questions detected")
        if not answers:
            warnings.append("No answer key found")
            p(f"  ⚠ No answer key detected")
    
    return issues, warnings


def audit_module_order(doc, name):
    """Part 6: Module logical order and objective matching."""
    p(f"\n{'─'*50}")
    p("PART 6: MODULE ORDER & OBJECTIVE MATCHING")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    
    # Extract Heading 1 entries (modules)
    modules = []
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'Heading 1':
            modules.append((i, para.text.strip()))
    
    p(f"  Total Heading 1 entries: {len(modules)}")
    p(f"\n  Module order:")
    
    core1_found = False
    core2_found = False
    current_section = 'INTRO'
    
    for idx, (para_i, title) in enumerate(modules):
        # Detect core section markers
        if 'CORE 1' in title.upper() or '220-1201' in title:
            core1_found = True
            current_section = 'CORE 1'
            p(f"\n  === CORE 1 (220-1201) ===")
        elif 'CORE 2' in title.upper() or '220-1202' in title:
            core2_found = True
            current_section = 'CORE 2'
            p(f"\n  === CORE 2 (220-1202) ===")
        
        p(f"    {idx+1}. {title}")
    
    if not core1_found:
        warnings.append("No explicit CORE 1 section marker found")
        p(f"\n  ⚠ No 'CORE 1' section divider found")
    if not core2_found:
        warnings.append("No explicit CORE 2 section marker found")
        p(f"\n  ⚠ No 'CORE 2' section divider found")
    
    # Check that objectives appear after module headings
    p(f"\n  Checking each module has objectives:")
    modules_without_objectives = []
    for idx, (para_i, title) in enumerate(modules):
        if 'Module' not in title and 'CORE' not in title.upper():
            continue
        
        # Look in next 5 paragraphs for "Module Objectives" or "Objectives"
        has_objectives = False
        for j in range(para_i + 1, min(para_i + 8, len(doc.paragraphs))):
            if 'Objective' in doc.paragraphs[j].text:
                has_objectives = True
                break
        
        if not has_objectives and 'Module' in title:
            modules_without_objectives.append(title[:50])
    
    if modules_without_objectives:
        warnings.append(f"{len(modules_without_objectives)} modules missing objectives")
        p(f"  ⚠ Modules without objectives section:")
        for m in modules_without_objectives:
            p(f"      - {m}")
    else:
        p(f"  ✓ All modules have objectives")
    
    # Check expected CompTIA A+ v15 topic coverage
    p(f"\n  Checking required topic coverage:")
    required_topics = [
        'Mobile Devices', 'Laptop', 'Networking',
        'Virtualization', 'Cloud', 'Hardware', 'Display',
        'Storage', 'Printer', 'Cable', 'Connector',
        'Operating System', 'Windows', 'Linux', 'macOS',
        'Security', 'Malware', 'Troubleshooting',
        'IoT', 'Embedded', 'Scripting',
        'Backup', 'Recovery', 'RAID',
        'Remote Access', 'MDM', 'Mobile Device Management',
        'Data Destruction', 'Licensing', 'AI', 'Artificial Intelligence',
    ]
    
    all_text = '\n'.join(para.text for para in doc.paragraphs).lower()
    missing_topics = []
    for topic in required_topics:
        if topic.lower() not in all_text:
            missing_topics.append(topic)
    
    if missing_topics:
        issues.append(f"{len(missing_topics)} required topics not found")
        p(f"  ❌ Missing topics: {missing_topics}")
    else:
        p(f"  ✓ All {len(required_topics)} required topics present")
    
    return issues, warnings


def audit_emoji_and_callouts(doc, name):
    """Check emoji removal and callout standardization."""
    p(f"\n{'─'*50}")
    p("PART 7: EMOJI & CALLOUTS")
    p(f"{'─'*50}")
    
    issues = []
    warnings = []
    
    # Emoji in headings
    heading_emoji = []
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith('Heading'):
            for char in para.text:
                if is_emoji(char):
                    heading_emoji.append((para.style.name, para.text[:50]))
                    break
    
    if heading_emoji:
        # Check if they're all ⚠️ variation selectors (minor)
        p(f"  ⚠ {len(heading_emoji)} headings with emoji/variation selectors")
        for style, text in heading_emoji[:5]:
            p(f"      [{style}] {text}")
        warnings.append(f"{len(heading_emoji)} headings with emoji remnants")
    else:
        p(f"  ✓ No emoji in headings")
    
    # Callout check
    callout_issues = []
    callout_counts = {'Tip -Trap': 0, 'Tip -Exam': 0, 'Tip -Pro': 0}
    
    for para in doc.paragraphs:
        if not para.style:
            continue
        style = para.style.name
        text = para.text.strip()
        
        if style == 'Tip -Trap':
            callout_counts['Tip -Trap'] += 1
            if not text.startswith('\u26A0'):
                callout_issues.append(f"Trap missing ⚠️: {text[:50]}")
        elif style == 'Tip -Exam':
            callout_counts['Tip -Exam'] += 1
            if not text.startswith('\u2139'):
                callout_issues.append(f"Exam missing ℹ️: {text[:50]}")
        elif style == 'Tip -Pro':
            callout_counts['Tip -Pro'] += 1
            if not text.startswith('\u2714'):
                callout_issues.append(f"Pro missing ✔: {text[:50]}")
    
    p(f"  Callout counts: Trap={callout_counts['Tip -Trap']}, Exam={callout_counts['Tip -Exam']}, Pro={callout_counts['Tip -Pro']}")
    
    if callout_issues:
        issues.append(f"{len(callout_issues)} callouts not standardized")
        p(f"  ❌ {len(callout_issues)} callouts not standardized:")
        for ci in callout_issues[:5]:
            p(f"      {ci}")
    else:
        p(f"  ✓ All callouts standardized")
    
    # Glossary bullets
    bullet_colors = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if '\u25cf' in run.text:
                try:
                    color = str(run.font.color.rgb) if run.font.color and run.font.color.rgb else 'none'
                except:
                    color = 'error'
                bullet_colors[color] = bullet_colors.get(color, 0) + 1
    
    if bullet_colors:
        p(f"  Glossary bullets: {bullet_colors}")
        if 'FF0000' in bullet_colors and 'FFC000' in bullet_colors and '00B300' in bullet_colors:
            p(f"  ✓ All 3 priority colors present")
        else:
            warnings.append("Glossary bullet colors incomplete")
    else:
        warnings.append("No glossary bullets found")
        p(f"  ⚠ No glossary bullets detected")
    
    return issues, warnings


def main():
    global out
    out = open(REPORT, 'w', encoding='utf-8')
    
    p("╔══════════════════════════════════════════════════════════════════════╗")
    p("║         COMPLETE PUBLISHING AUDIT REPORT                           ║")
    p("║         CompTIA A+ 220-1200 Series Study Guide                     ║")
    p("╚══════════════════════════════════════════════════════════════════════╝")
    
    all_issues = {}
    all_warnings = {}
    
    for doc_name, filepath in FILES.items():
        try:
            doc = Document(filepath)
        except Exception as e:
            p(f"\n❌ ERROR loading {filepath}: {e}")
            continue
        
        all_issues[doc_name] = []
        all_warnings[doc_name] = []
        
        # Part 1: Structure
        i, w = audit_document(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 2: Spelling/merging
        i, w = audit_spelling_and_merging(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 3: Images
        i, w = audit_images(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 4: Content accuracy
        i, w = audit_content_accuracy(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 5: Q&A
        i, w = audit_questions_answers(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 6: Module order
        i, w = audit_module_order(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
        
        # Part 7: Emoji & callouts
        i, w = audit_emoji_and_callouts(doc, doc_name)
        all_issues[doc_name].extend(i)
        all_warnings[doc_name].extend(w)
    
    # === FINAL SUMMARY ===
    p(f"\n\n{'═'*70}")
    p("FINAL AUDIT SUMMARY & PUBLISH READINESS")
    p(f"{'═'*70}")
    
    for doc_name in FILES.keys():
        if doc_name not in all_issues:
            continue
        p(f"\n{'─'*40}")
        p(f"  {doc_name}")
        p(f"{'─'*40}")
        
        issue_count = len(all_issues[doc_name])
        warn_count = len(all_warnings[doc_name])
        
        p(f"  Critical Issues: {issue_count}")
        for issue in all_issues[doc_name]:
            p(f"    ❌ {issue}")
        
        p(f"\n  Warnings: {warn_count}")
        for w in all_warnings[doc_name]:
            p(f"    ⚠ {w}")
        
        # Publish readiness verdict
        p(f"\n  PUBLISH READINESS:")
        if issue_count == 0:
            p(f"    ✅ READY FOR PUBLISHING")
            p(f"    Document passes all critical checks.")
        elif issue_count <= 3:
            p(f"    🟡 NEARLY READY - {issue_count} minor issues to address")
            p(f"    Fix the issues above before final submission.")
        else:
            p(f"    🔴 NOT READY - {issue_count} issues require attention")
            p(f"    Significant corrections needed before publishing.")
    
    p(f"\n{'═'*70}")
    p("END OF AUDIT")
    p(f"{'═'*70}")
    
    out.close()
    print(f"Complete audit written to {REPORT}")


if __name__ == '__main__':
    main()
