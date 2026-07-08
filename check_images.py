"""Quick check: where did the images go?"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

out = open('image_check.txt', 'w', encoding='utf-8')

# Check original
doc = Document(r'public\reference\book\comptia-book.docx')
out.write("=== ORIGINAL DOCUMENT IMAGES ===\n")
count = 0
for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        count += 1
        style = para.style.name if para.style else 'None'
        text = para.text[:60]
        out.write(f"  Para {i} [{style}]: {len(drawings)} drawing(s) | text: '{text}'\n")

out.write(f"\nTotal paragraphs with images: {count}\n")

# Also check document body XML for total drawings
body = doc.element.body
all_drawings = body.findall('.//' + qn('w:drawing'))
out.write(f"Total w:drawing elements in body: {len(all_drawings)}\n")

# Check which are in TOC-styled paragraphs
toc_images = 0
for para in doc.paragraphs:
    if para.style and para.style.name.lower().startswith('toc'):
        drawings = para._element.findall('.//' + qn('w:drawing'))
        if drawings:
            toc_images += len(drawings)
            out.write(f"  TOC image: [{para.style.name}] {para.text[:50]}\n")

out.write(f"\nImages in TOC paragraphs: {toc_images}\n")

# Check output
out.write("\n\n=== OUTPUT DOCUMENT IMAGES ===\n")
doc2 = Document(r'public\reference\book\comptia-book-KDP-v3.docx')
count2 = 0
for i, para in enumerate(doc2.paragraphs):
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        count2 += 1
        style = para.style.name if para.style else 'None'
        text = para.text[:60]
        out.write(f"  Para {i} [{style}]: {len(drawings)} drawing(s) | text: '{text}'\n")

out.write(f"\nTotal paragraphs with images: {count2}\n")

body2 = doc2.element.body
all_drawings2 = body2.findall('.//' + qn('w:drawing'))
out.write(f"Total w:drawing elements in body: {len(all_drawings2)}\n")

out.write(f"\nDIFFERENCE: {count - count2} paragraphs lost images\n")
out.close()
print("Done - see image_check.txt")
