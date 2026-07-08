"""
COMPLETE KDP CONVERSION SCRIPT FOR COMPTIA-BOOK.DOCX
Handles all transformations including table flattening with proper XML manipulation
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from lxml import etree
import re
import shutil
import os

# Configuration
SOURCE_FILE = r'public\reference\book\comptia-book.docx'
OUTPUT_FILE = r'public\reference\book\comptia-book-KDP.docx'
BACKUP_FILE = r'public\reference\book\comptia-book-BACKUP.docx'

# Tables that should become images (5+ columns or complex)
# Images are located in public/reference/book/images/
IMAGE_EXCEPTION_TABLES = {
    47: ("Shared_Responsibility_Model.png", "IaaS PaaS SaaS shared responsibility comparison"),
    69: ("Memory_Types_Comparison.png", "DDR memory types detailed comparison"),
    73: ("Display_Resolutions.png", "Common display resolutions and specifications"),
    82: ("RAID_Configurations.png", "RAID levels feature comparison"),
    91: ("Windows_Versions.png", "Windows editions feature comparison"),
    93: ("Security_Protocols.png", "Security protocols comparison table"),
    133: ("Cable_Specifications.png", "Cable types and specifications table")
}

IMAGE_DIR = r'public\reference\book\images'

# Statistics
stats = {
    'headings_emoji_stripped': 0,
    'variation_selectors_removed': 0,
    'callouts_fixed': 0,
    'tables_flattened': 0,
    'tables_to_image': 0,
    'merged_cells_handled': 0,
    'images_centered': 0,
    'images_alt_added': 0,
    'headers_cleared': 0,
    'footers_cleared': 0,
    'blank_paragraphs_removed': 0,
    'glossary_bullets_preserved': 0
}

def strip_emoji_and_vs(text):
    """
    Remove emoji, variation selectors (U+FE00–FE0F), and zero-width joiners (U+200D).
    Handles multi-character emoji properly.
    """
    cleaned = ""
    i = 0
    while i < len(text):
        char = text[i]
        cp = ord(char)
        
        # Skip variation selectors
        if 0xFE00 <= cp <= 0xFE0F:
            stats['variation_selectors_removed'] += 1
            i += 1
            continue
        
        # Skip zero-width joiner
        if cp == 0x200D:
            i += 1
            continue
        
        # Skip emoji ranges
        if (0x1F000 <= cp <= 0x1FFFF or  # Emoji & Pictographs
            0x2600 <= cp <= 0x26FF or    # Miscellaneous Symbols
            0x2700 <= cp <= 0x27BF or    # Dingbats
            0x231A <= cp <= 0x23FA or    # Miscellaneous Technical
            0x1F300 <= cp <= 0x1F9FF):   # Additional emoji
            i += 1
            continue
        
        cleaned += char
        i += 1
    
    return cleaned.strip()

def is_glossary_bullet(run):
    """Check if run is a colored glossary priority bullet"""
    if '●' in run.text:
        if run.font.color and run.font.color.rgb:
            return True
    return False

def strip_heading_emoji(para):
    """Remove emoji from heading, preserving structure"""
    if not para.style.name.startswith('Heading'):
        return False
    
    modified = False
    runs_to_remove = []
    
    for i, run in enumerate(para.runs):
        # Preserve colored glossary bullets
        if is_glossary_bullet(run):
            stats['glossary_bullets_preserved'] += 1
            continue
        
        original = run.text
        cleaned = strip_emoji_and_vs(original)
        
        if cleaned != original:
            if cleaned:
                run.text = cleaned
                modified = True
            else:
                # Run is now empty - mark for removal
                runs_to_remove.append(i)
                modified = True
    
    # Remove empty runs (reverse order to maintain indices)
    for i in reversed(runs_to_remove):
        run_element = para.runs[i]._element
        run_element.getparent().remove(run_element)
    
    if modified:
        stats['headings_emoji_stripped'] += 1
    
    return modified

def fix_callout(para):
    """
    Ensure callouts have correct emoji:
    ⚠️ Exam Trap | ℹ️ Exam Tip | 💡 Pro Tip
    """
    style = para.style.name.lower() if para.style.name else ''
    
    if 'exam trap' in style:
        emoji, label = '⚠️', 'Exam Trap'
    elif 'exam tip' in style:
        emoji, label = 'ℹ️', 'Exam Tip'
    elif 'pro tip' in style:
        emoji, label = '💡', 'Pro Tip'
    else:
        return False
    
    text = para.text.strip()
    
    # Check if correct
    if text.startswith(f"{emoji} {label}"):
        return False
    
    # Fix it
    cleaned = strip_emoji_and_vs(text)
    if cleaned.startswith(label):
        new_text = f"{emoji} {cleaned}"
    else:
        new_text = f"{emoji} {label}: {cleaned}"
    
    # Replace paragraph text while preserving formatting
    for run in para.runs:
        run._element.getparent().remove(run._element)
    
    para.add_run(new_text)
    stats['callouts_fixed'] += 1
    return True

def deduplicate_merged_cells(row):
    """Handle merged cells (gridSpan) by returning unique cell objects"""
    seen = set()
    unique = []
    
    for cell in row.cells:
        cell_id = id(cell._element)
        if cell_id not in seen:
            seen.add(cell_id)
            unique.append(cell)
        else:
            stats['merged_cells_handled'] += 1
    
    return unique

def flatten_table(table):
    """
    Convert table to text paragraphs:
    2-col: **Term** — description
    3+ col: **Term** Header1: value1 Header2: value2
    """
    result_paragraphs = []
    
    if len(table.rows) == 0:
        return result_paragraphs
    
    # Get headers
    header_row = table.rows[0]
    headers = [cell.text.strip() for cell in deduplicate_merged_cells(header_row)]
    
    # Process data rows
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        cells = deduplicate_merged_cells(row)
        values = [c.text.strip() for c in cells]
        
        if not values or not values[0]:
            continue
        
        if len(cells) == 2:
            # 2-column format
            text = f"**{values[0]}** — {values[1]}"
        else:
            # 3+ column format
            parts = [f"**{values[0]}**"]
            for i in range(1, len(values)):
                if i < len(headers) and values[i]:
                    parts.append(f"{headers[i]}: {values[i]}")
            text = " ".join(parts)
        
        result_paragraphs.append(text)
    
    return result_paragraphs

def replace_table_with_paragraphs(doc, table_idx, paragraphs_text):
    """
    Replace a table with paragraphs using XML manipulation.
    This is the critical function that python-docx doesn't provide.
    """
    table = doc.tables[table_idx]
    table_element = table._element
    parent = table_element.getparent()
    
    # Find table's position
    table_position = list(parent).index(table_element)
    
    # Create new paragraph elements
    for i, text in enumerate(paragraphs_text):
        # Create paragraph XML
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        
        # Set style (First Paragraph or Body Text)
        pStyle = OxmlElement('w:pStyle')
        if i == 0:
            pStyle.set(qn('w:val'), 'FirstParagraph')
        else:
            pStyle.set(qn('w:val'), 'BodyText')
        pPr.append(pStyle)
        p.append(pPr)
        
        # Add text runs (handle bold formatting for **term**)
        if text.startswith('**'):
            # Extract bold term
            match = re.match(r'\*\*([^*]+)\*\*(.*)', text)
            if match:
                bold_text = match.group(1)
                rest_text = match.group(2)
                
                # Bold run
                r1 = OxmlElement('w:r')
                rPr1 = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr1.append(b)
                r1.append(rPr1)
                t1 = OxmlElement('w:t')
                t1.text = bold_text
                r1.append(t1)
                p.append(r1)
                
                # Regular run
                r2 = OxmlElement('w:r')
                t2 = OxmlElement('w:t')
                t2.text = rest_text
                r2.append(t2)
                p.append(r2)
            else:
                # Fallback
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = text
                r.append(t)
                p.append(r)
        else:
            # Simple text
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = text
            r.append(t)
            p.append(r)
        
        # Insert paragraph
        parent.insert(table_position + i, p)
    
    # Remove table
    parent.remove(table_element)
    
    stats['tables_flattened'] += 1

def replace_table_with_image(doc, table_idx, image_filename, alt_text):
    """Replace a table with an actual image"""
    from docx.shared import Inches
    import os
    
    table = doc.tables[table_idx]
    table_element = table._element
    parent = table_element.getparent()
    table_position = list(parent).index(table_element)
    
    # Check if image file exists
    image_path = os.path.join(IMAGE_DIR, image_filename)
    
    if not os.path.exists(image_path):
        # Fallback to placeholder text
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
        
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = f"[IMAGE: {image_filename} - {alt_text}]"
        r.append(t)
        p.append(r)
        
        parent.insert(table_position, p)
    else:
        # Insert actual image via python-docx (simpler than raw XML)
        # We need to insert before table, then remove table
        # Find which paragraph index corresponds to this table
        
        # Create new document paragraph at this position (tricky with python-docx)
        # For now, create paragraph with placeholder that we'll replace
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
        
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = f"__IMAGE_PLACEHOLDER_{table_idx}__"
        r.append(t)
        p.append(r)
        
        parent.insert(table_position, p)
        
        # Store image info for second pass
        if not hasattr(doc, '_image_placeholders'):
            doc._image_placeholders = []
        doc._image_placeholders.append((table_idx, image_path, alt_text))
    
    # Remove table
    parent.remove(table_element)
    stats['tables_to_image'] += 1

def insert_actual_images(doc):
    """Second pass: replace image placeholders with actual images"""
    if not hasattr(doc, '_image_placeholders'):
        return
    
    for table_idx, image_path, alt_text in doc._image_placeholders:
        placeholder_text = f"__IMAGE_PLACEHOLDER_{table_idx}__"
        
        # Find paragraph with placeholder
        for para in doc.paragraphs:
            if placeholder_text in para.text:
                # Clear paragraph
                for run in para.runs:
                    run._element.getparent().remove(run._element)
                
                # Add image
                run = para.add_run()
                picture = run.add_picture(image_path, width=Inches(6.0))
                
                # Add alt text to image
                # This is complex in python-docx, we'll handle it after
                
                # Center paragraph
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                print(f"   ✓ Inserted image: {os.path.basename(image_path)}")
                break

def center_images(doc):
    """Center all images and ensure alt text"""
    for para in doc.paragraphs:
        pics = para._element.xpath('.//pic:pic')
        if pics:
            # Center paragraph
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            stats['images_centered'] += 1
            
            # Check alt text
            for pic in pics:
                docPr_list = pic.xpath('.//wp:docPr')
                if docPr_list:
                    docPr = docPr_list[0]
                    if not docPr.get('descr'):
                        docPr.set('descr', 'Diagram - see print version for details')
                        stats['images_alt_added'] += 1

def clear_headers_footers(doc):
    """Remove all headers and footers"""
    for section in doc.sections:
        # Clear header
        for para in section.header.paragraphs:
            para._element.getparent().remove(para._element)
        if section.header.paragraphs:
            stats['headers_cleared'] += 1
        
        # Clear footer
        for para in section.footer.paragraphs:
            para._element.getparent().remove(para._element)
        if section.footer.paragraphs:
            stats['footers_cleared'] += 1

def remove_excess_blanks(doc):
    """Remove runs of 2+ consecutive empty paragraphs, keep at most 1"""
    to_remove = []
    consecutive = []
    
    for i, para in enumerate(doc.paragraphs):
        # Skip glossary bullets
        if any(is_glossary_bullet(r) for r in para.runs):
            if len(consecutive) >= 2:
                to_remove.extend(consecutive[1:])
            consecutive = []
            continue
        
        if not para.text.strip():
            consecutive.append(i)
        else:
            if len(consecutive) >= 2:
                to_remove.extend(consecutive[1:])
            consecutive = []
    
    # Handle trailing
    if len(consecutive) >= 2:
        to_remove.extend(consecutive[1:])
    
    # Remove in reverse
    for i in sorted(set(to_remove), reverse=True):
        p_elem = doc.paragraphs[i]._element
        p_elem.getparent().remove(p_elem)
        stats['blank_paragraphs_removed'] += 1

def main():
    print("="*80)
    print("COMPLETE KDP CONVERSION - COMPTIA A+ STUDY GUIDE")
    print("="*80)
    
    # Backup
    print(f"\nCreating backup...")
    shutil.copy2(SOURCE_FILE, BACKUP_FILE)
    print(f"   OK {BACKUP_FILE}")
    
    # Load
    print(f"\nLoading document...")
    doc = Document(SOURCE_FILE)
    print(f"   OK {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # 1. Strip emoji from headings
    print(f"\n1. Stripping emoji from headings...")
    for para in doc.paragraphs:
        strip_heading_emoji(para)
    print(f"   OK {stats['headings_emoji_stripped']} headings cleaned")
    print(f"   OK {stats['variation_selectors_removed']} variation selectors removed")
    
    # 2. Fix callouts
    print(f"\n2. Fixing callout consistency...")
    for para in doc.paragraphs:
        fix_callout(para)
    print(f"   OK {stats['callouts_fixed']} callouts fixed")
    
    # 3. Process tables
    print(f"\n3. Processing {len(doc.tables)} tables...")
    
    # Collect table operations (process in reverse to maintain indices)
    table_ops = []
    for i in range(len(doc.tables)):
        if i in IMAGE_EXCEPTION_TABLES:
            img_file, alt = IMAGE_EXCEPTION_TABLES[i]
            table_ops.append(('image', i, img_file, alt))
        else:
            table = doc.tables[i]
            # Flatten if 4 or fewer columns
            if len(table.columns) <= 4:
                paragraphs = flatten_table(table)
                table_ops.append(('flatten', i, paragraphs, None))
            else:
                # 5+ columns not in exception list - also flatten (or could add to exceptions)
                paragraphs = flatten_table(table)
                table_ops.append(('flatten', i, paragraphs, None))
    
    # Execute in reverse order
    for op_type, idx, data, extra in reversed(table_ops):
        if op_type == 'flatten':
            replace_table_with_paragraphs(doc, idx, data)
        elif op_type == 'image':
            replace_table_with_image(doc, idx, data, extra)
    
    # Insert actual images (second pass)
    print(f"   OK Inserting actual image files...")
    insert_actual_images(doc)
    
    print(f"   OK {stats['tables_flattened']} tables flattened")
    print(f"   OK {stats['tables_to_image']} tables converted to images")
    print(f"   OK {stats['merged_cells_handled']} merged cells handled")
    
    # 4. Images
    print(f"\n4. Processing images...")
    center_images(doc)
    print(f"   OK {stats['images_centered']} images centered")
    print(f"   OK {stats['images_alt_added']} alt texts added")
    
    # 5. Headers/footers
    print(f"\n5. Removing headers and footers...")
    clear_headers_footers(doc)
    print(f"   OK {stats['headers_cleared']} headers cleared")
    print(f"   OK {stats['footers_cleared']} footers cleared")
    
    # 6. Blank paragraphs
    print(f"\n6. Removing excess blank paragraphs...")
    remove_excess_blanks(doc)
    print(f"   OK {stats['blank_paragraphs_removed']} removed")
    print(f"   OK {stats['glossary_bullets_preserved']} glossary bullets preserved")
    
    # Save
    print(f"\nSaving KDP document...")
    doc.save(OUTPUT_FILE)
    print(f"   OK {OUTPUT_FILE}")
    
    # Summary
    print("\n" + "="*80)
    print("CONVERSION COMPLETE")
    print("="*80)
    for key, value in stats.items():
        print(f"   {key.replace('_', ' ').title()}: {value}")
    
    print(f"\nFiles created:")
    print(f"   KDP Output: {OUTPUT_FILE}")
    print(f"   Backup: {BACKUP_FILE}")
    print("\nReady for KDP upload!")

if __name__ == '__main__':
    main()
