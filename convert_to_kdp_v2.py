"""
REVISED KDP CONVERSION SCRIPT - Version 2
Changes:
- Flatten ALL tables to text (no images)
- Remove Table of Contents
- Preserve existing alt text on images
- Remove all page breaks including from headings
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from lxml import etree
import re
import shutil
import os

# Configuration
SOURCE_FILE = r'public\reference\book\comptia-book.docx'
OUTPUT_FILE = r'public\reference\book\comptia-book-KDP-v2.docx'
BACKUP_FILE = r'public\reference\book\comptia-book-BACKUP-v2.docx'

# Statistics
stats = {
    'headings_emoji_stripped': 0,
    'variation_selectors_removed': 0,
    'callouts_fixed': 0,
    'tables_flattened': 0,
    'merged_cells_handled': 0,
    'images_centered': 0,
    'images_alt_preserved': 0,
    'headers_cleared': 0,
    'footers_cleared': 0,
    'blank_paragraphs_removed': 0,
    'glossary_bullets_preserved': 0,
    'page_breaks_removed': 0,
    'toc_removed': 0
}

# Emoji ranges to strip
EMOJI_PATTERN = re.compile(
    '[\U0001F000-\U0001FFFF]|'  # Emoji symbols
    '[\U00002600-\U000026FF]|'  # Misc symbols
    '[\U00002700-\U000027BF]|'  # Dingbats
    '[\U0000FE00-\U0000FE0F]|'  # Variation selectors
    '[\U0000200D]'              # Zero-width joiner
)

def strip_heading_emoji(para):
    """Remove emoji from heading paragraphs"""
    if not para.style.name.startswith('Heading'):
        return
    
    for run in para.runs[:]:
        original = run.text
        cleaned = EMOJI_PATTERN.sub('', original)
        
        if original != cleaned:
            stats['headings_emoji_stripped'] += 1
            run.text = cleaned
            
            if '\uFE0F' in original or '\u200D' in original:
                stats['variation_selectors_removed'] += 1
        
        # Remove empty runs
        if not run.text.strip():
            run._element.getparent().remove(run._element)

def remove_page_breaks(para):
    """Remove page breaks from paragraphs including headings"""
    # Remove page breaks from paragraph properties
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        # Remove pageBreakBefore
        pageBreakBefore = pPr.find(qn('w:pageBreakBefore'))
        if pageBreakBefore is not None:
            pPr.remove(pageBreakBefore)
            stats['page_breaks_removed'] += 1
    
    # Remove explicit page break runs
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            br_type = br.get(qn('w:type'))
            if br_type == 'page':
                run._element.remove(br)
                stats['page_breaks_removed'] += 1

def fix_callout(para):
    """Fix callout emoji consistency"""
    style = para.style.name if para.style else ''
    
    callout_map = {
        'Exam Trap': '⚠️',
        'Exam Tip': 'ℹ️',
        'Pro Tip': '💡'
    }
    
    for callout_name, correct_emoji in callout_map.items():
        if callout_name.lower() in style.lower():
            if para.runs:
                first_run = para.runs[0]
                if not first_run.text.startswith(correct_emoji):
                    # Strip any existing emoji
                    first_run.text = EMOJI_PATTERN.sub('', first_run.text).lstrip()
                    # Add correct emoji
                    first_run.text = f"{correct_emoji} {first_run.text}"
                    stats['callouts_fixed'] += 1
            break

def is_glossary_bullet(run):
    """Check if run is a colored glossary bullet"""
    if '●' in run.text:
        try:
            if run.font.color and run.font.color.rgb:
                return True
        except:
            pass
    return False

def deduplicate_merged_cells(row):
    """Handle merged cells by deduplicating based on cell element ID"""
    seen = set()
    unique = []
    for cell in row.cells:
        cell_id = id(cell._element)
        if cell_id not in seen:
            seen.add(cell_id)
            unique.append(cell)
    return unique

def flatten_table(table):
    """Convert table to formatted text paragraphs"""
    paragraphs = []
    num_cols = len(table.columns)
    
    # Check for merged cells
    has_merged = False
    for row in table.rows:
        unique_cells = deduplicate_merged_cells(row)
        if len(unique_cells) < len(row.cells):
            has_merged = True
            stats['merged_cells_handled'] += 1
            break
    
    for row_idx, row in enumerate(table.rows):
        cells = deduplicate_merged_cells(row)
        cell_texts = [cell.text.strip() for cell in cells]
        
        # Skip completely empty rows
        if not any(cell_texts):
            continue
        
        # Header row or first column as term
        if row_idx == 0 or num_cols == 2:
            if num_cols == 2 and len(cell_texts) >= 2:
                # 2-column: **Term** — description
                term = cell_texts[0]
                desc = cell_texts[1]
                text = f"**{term}** — {desc}"
                paragraphs.append(('Body Text', text))
            elif len(cell_texts) >= 1:
                # Multi-column: **Term** Header1: value1, Header2: value2
                term = cell_texts[0]
                parts = [f"**{term}**"]
                
                # If we have a header row, use it
                if row_idx == 0:
                    # This is the header - just bold all items
                    text = " | ".join([f"**{cell}**" for cell in cell_texts])
                    paragraphs.append(('Body Text', text))
                    continue
                
                # Data rows - format as term with attributes
                for i, value in enumerate(cell_texts[1:], 1):
                    if value:
                        # Try to get header from first row if available
                        header = f"Col{i}"
                        if len(table.rows) > 0:
                            header_cells = deduplicate_merged_cells(table.rows[0])
                            if i < len(header_cells):
                                header = header_cells[i].text.strip() or f"Col{i}"
                        parts.append(f"{header}: {value}")
                
                text = " ".join(parts)
                paragraphs.append(('Body Text', text))
        else:
            # Subsequent rows
            if num_cols == 2 and len(cell_texts) >= 2:
                term = cell_texts[0]
                desc = cell_texts[1]
                text = f"**{term}** — {desc}"
                paragraphs.append(('Body Text', text))
            elif len(cell_texts) >= 1:
                term = cell_texts[0]
                parts = [f"**{term}**"]
                
                # Get headers from first row
                header_cells = deduplicate_merged_cells(table.rows[0])
                header_texts = [c.text.strip() for c in header_cells]
                
                for i, value in enumerate(cell_texts[1:], 1):
                    if value:
                        header = header_texts[i] if i < len(header_texts) else f"Col{i}"
                        parts.append(f"{header}: {value}")
                
                text = " ".join(parts)
                paragraphs.append(('Body Text', text))
    
    stats['tables_flattened'] += 1
    return paragraphs

def replace_table_with_paragraphs(doc, table_idx, paragraphs):
    """Replace table with text paragraphs using XML manipulation"""
    table = doc.tables[table_idx]
    table_element = table._element
    parent = table_element.getparent()
    table_position = list(parent).index(table_element)
    
    # Insert paragraphs before table
    for style_name, text in paragraphs:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_name)
        pPr.append(pStyle)
        p.append(pPr)
        
        # Parse text for bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                b = OxmlElement('w:b')
                rPr.append(b)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.text = part[2:-2]
                r.append(t)
                p.append(r)
            elif part:
                # Normal text
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = part
                r.append(t)
                p.append(r)
        
        parent.insert(table_position, p)
        table_position += 1
    
    # Remove table
    parent.remove(table_element)

def center_images_preserve_alt(doc):
    """Center all images and preserve existing alt text"""
    for para in doc.paragraphs:
        # Check if paragraph contains images
        pics = para._element.xpath('.//pic:pic')
        
        if pics:
            # Center the paragraph
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            stats['images_centered'] += 1
            
            # Check and preserve alt text
            for pic in pics:
                docPr_list = pic.xpath('.//wp:docPr')
                if docPr_list:
                    docPr = docPr_list[0]
                    existing_alt = docPr.get('descr')
                    if existing_alt:
                        stats['images_alt_preserved'] += 1
                    # If no alt text exists, we could add a default, but instruction says preserve
                    # So we only count preserved ones

def remove_toc(doc):
    """Remove Table of Contents"""
    # TOC is usually in a field or has a specific style
    # Look for TOC fields and TOC styles
    
    paragraphs_to_remove = []
    
    for i, para in enumerate(doc.paragraphs):
        # Check for TOC field
        fldSimple = para._element.find(qn('w:fldSimple'))
        if fldSimple is not None:
            instr = fldSimple.get(qn('w:instr'), '')
            if 'TOC' in instr:
                paragraphs_to_remove.append(i)
                continue
        
        # Check for TOC styles
        if para.style and 'toc' in para.style.name.lower():
            paragraphs_to_remove.append(i)
            continue
        
        # Check for hyperlinks in runs (TOC entries are often hyperlinks)
        # But be careful - not all hyperlinks are TOC
        # Look for patterns like "Module 1" followed by page numbers
        if para.text.strip() and re.search(r'(Module|Chapter|Section)\s+\d+.*\d+$', para.text):
            # Could be TOC entry
            # Additional check: does it have a hyperlink?
            hyperlinks = para._element.xpath('.//w:hyperlink')
            if hyperlinks:
                paragraphs_to_remove.append(i)
    
    # Also check for "Table of Contents" or "Contents" heading
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name.startswith('Heading'):
            text = para.text.strip().lower()
            if text in ['table of contents', 'contents', 'toc']:
                paragraphs_to_remove.append(i)
                # Remove next N paragraphs that look like TOC entries
                for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    # Stop at next heading
                    if next_para.style and next_para.style.name.startswith('Heading'):
                        break
                    # If it looks like a TOC entry, mark for removal
                    if re.search(r'\d+$', next_para.text.strip()):
                        paragraphs_to_remove.append(j)
    
    # Remove paragraphs (in reverse order to maintain indices)
    for idx in sorted(set(paragraphs_to_remove), reverse=True):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para_element = para._element
            para_element.getparent().remove(para_element)
            stats['toc_removed'] += 1

def clear_headers_footers(doc):
    """Remove all headers and footers"""
    for section in doc.sections:
        # Headers
        for para in section.header.paragraphs:
            if para.text.strip():
                stats['headers_cleared'] += 1
            para.clear()
        
        # Footers
        for para in section.footer.paragraphs:
            if para.text.strip():
                stats['footers_cleared'] += 1
            para.clear()

def remove_excess_blanks(doc):
    """Remove consecutive empty paragraphs, keep max 1"""
    consecutive_empty = []
    to_remove = []
    
    for i, para in enumerate(doc.paragraphs):
        # Skip glossary bullet paragraphs
        if any(is_glossary_bullet(run) for run in para.runs):
            stats['glossary_bullets_preserved'] += 1
            consecutive_empty = []
            continue
        
        if not para.text.strip():
            consecutive_empty.append(i)
        else:
            # Process accumulated empty paragraphs
            if len(consecutive_empty) >= 2:
                # Remove all but first
                to_remove.extend(consecutive_empty[1:])
            consecutive_empty = []
    
    # Handle trailing empties
    if len(consecutive_empty) >= 2:
        to_remove.extend(consecutive_empty[1:])
    
    # Remove in reverse order
    for idx in sorted(to_remove, reverse=True):
        para = doc.paragraphs[idx]
        para_element = para._element
        para_element.getparent().remove(para_element)
        stats['blank_paragraphs_removed'] += 1

def main():
    print("="*80)
    print("REVISED KDP CONVERSION - VERSION 2")
    print("="*80)
    
    # Backup
    print(f"\nCreating backup...")
    shutil.copy2(SOURCE_FILE, BACKUP_FILE)
    print(f"   OK {BACKUP_FILE}")
    
    # Load
    print(f"\nLoading document...")
    doc = Document(SOURCE_FILE)
    print(f"   OK {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # 1. Remove TOC
    print(f"\n1. Removing Table of Contents...")
    remove_toc(doc)
    print(f"   OK {stats['toc_removed']} TOC entries removed")
    
    # 2. Strip emoji from headings
    print(f"\n2. Stripping emoji from headings...")
    for para in doc.paragraphs:
        strip_heading_emoji(para)
    print(f"   OK {stats['headings_emoji_stripped']} headings cleaned")
    print(f"   OK {stats['variation_selectors_removed']} variation selectors removed")
    
    # 3. Remove page breaks
    print(f"\n3. Removing page breaks...")
    for para in doc.paragraphs:
        remove_page_breaks(para)
    print(f"   OK {stats['page_breaks_removed']} page breaks removed")
    
    # 4. Fix callouts
    print(f"\n4. Fixing callout consistency...")
    for para in doc.paragraphs:
        fix_callout(para)
    print(f"   OK {stats['callouts_fixed']} callouts fixed")
    
    # 5. Flatten ALL tables
    print(f"\n5. Flattening {len(doc.tables)} tables...")
    
    # Collect table operations (process in reverse to maintain indices)
    table_ops = []
    for i in range(len(doc.tables)):
        table = doc.tables[i]
        paragraphs = flatten_table(table)
        table_ops.append((i, paragraphs))
    
    # Execute in reverse order
    for idx, paragraphs in reversed(table_ops):
        replace_table_with_paragraphs(doc, idx, paragraphs)
    
    print(f"   OK {stats['tables_flattened']} tables flattened")
    print(f"   OK {stats['merged_cells_handled']} tables with merged cells handled")
    
    # 6. Center images and preserve alt text
    print(f"\n6. Processing images...")
    center_images_preserve_alt(doc)
    print(f"   OK {stats['images_centered']} images centered")
    print(f"   OK {stats['images_alt_preserved']} alt texts preserved")
    
    # 7. Headers/footers
    print(f"\n7. Removing headers and footers...")
    clear_headers_footers(doc)
    print(f"   OK {stats['headers_cleared']} headers cleared")
    print(f"   OK {stats['footers_cleared']} footers cleared")
    
    # 8. Blank paragraphs
    print(f"\n8. Removing excess blank paragraphs...")
    remove_excess_blanks(doc)
    print(f"   OK {stats['blank_paragraphs_removed']} removed")
    print(f"   OK {stats['glossary_bullets_preserved']} glossary bullets preserved")
    
    # Save
    print(f"\nSaving KDP document...")
    doc.save(OUTPUT_FILE)
    print(f"   OK {OUTPUT_FILE}")
    
    # Summary
    print("\n" + "="*80)
    print("CONVERSION COMPLETE")
    print("="*80)
    for key, value in stats.items():
        print(f"   {key.replace('_', ' ').title()}: {value}")
    
    print(f"\nFiles created:")
    print(f"   KDP Output: {OUTPUT_FILE}")
    print(f"   Backup: {BACKUP_FILE}")
    print("\nReady for KDP upload!")

if __name__ == '__main__':
    main()
