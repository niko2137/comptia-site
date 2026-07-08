"""Find the image missing alt text in KDP-final.docx"""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'public\reference\book\KDP-final.docx')
out = open('missing_alt_image.txt', 'w', encoding='utf-8')

for i, para in enumerate(doc.paragraphs):
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        for d in drawings:
            docPr = d.find('.//' + qn('wp:docPr'))
            if docPr is not None:
                alt = docPr.get('descr', '')
                name = docPr.get('name', '(unnamed)')
                if not alt:
                    # Get surrounding context
                    prev_text = doc.paragraphs[i-1].text[:80] if i > 0 else ''
                    next_text = doc.paragraphs[i+1].text[:80] if i < len(doc.paragraphs)-1 else ''
                    out.write(f"MISSING ALT TEXT:\n")
                    out.write(f"  Paragraph index: {i}\n")
                    out.write(f"  Image name: {name}\n")
                    out.write(f"  Para style: {para.style.name if para.style else 'None'}\n")
                    out.write(f"  Para text: '{para.text[:80]}'\n")
                    out.write(f"  Previous para: '{prev_text}'\n")
                    out.write(f"  Next para: '{next_text}'\n")
            else:
                out.write(f"  Para {i}: No docPr element at all\n")

out.close()
print("Done")
