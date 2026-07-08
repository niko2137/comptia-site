from docx import Document

book_path = r'public\reference\book\book.docx'

print("="*70)
print("VERIFICATION: Quality Improvements in book.docx")
print("="*70)

doc = Document(book_path)

# Check for each improvement
improvements = {
    "IPv6 Address Format and Notation": False,
    "SLAAC: Stateless Address Autoconfiguration": False,
    "Common IPv6 Prefix Types": False,
    "USB Power Delivery 3.1": False,
    "Wi-Fi 6E Deep Dive": False,
    "Thread Protocol": False,
    "Physical Safety & Environmental Controls": False,
    "Troubleshooting Methodology & Procedures": False
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if "IPv6 Address Format and Notation" in text:
        improvements["IPv6 Address Format and Notation"] = True
        print(f"\n✓ Found 'IPv6 Address Format and Notation' at paragraph {i}")
    
    if "SLAAC" in text and "Stateless" in text:
        improvements["SLAAC: Stateless Address Autoconfiguration"] = True
        print(f"✓ Found 'SLAAC' section at paragraph {i}")
    
    if "Common IPv6 Prefix" in text or "fe80::/10" in text:
        improvements["Common IPv6 Prefix Types"] = True
        print(f"✓ Found 'Common IPv6 Prefix Types' at paragraph {i}")
    
    if "USB Power Delivery 3.1" in text or ("USB PD 3.1" in text):
        improvements["USB Power Delivery 3.1"] = True
        print(f"✓ Found 'USB Power Delivery 3.1' at paragraph {i}")
    
    if "Wi-Fi 6E Deep Dive" in text:
        improvements["Wi-Fi 6E Deep Dive"] = True
        print(f"✓ Found 'Wi-Fi 6E Deep Dive' at paragraph {i}")
    
    if "Thread Protocol" in text and "🔗" in text:
        improvements["Thread Protocol"] = True
        print(f"✓ Found 'Thread Protocol' at paragraph {i}")
    
    if "Physical Safety & Environmental Controls" in text:
        improvements["Physical Safety & Environmental Controls"] = True
        print(f"✓ Found module rename 'Physical Safety & Environmental Controls' at paragraph {i}")
    
    if "Troubleshooting Methodology & Procedures" in text:
        improvements["Troubleshooting Methodology & Procedures"] = True
        print(f"✓ Found module rename 'Troubleshooting Methodology & Procedures' at paragraph {i}")

print("\n" + "="*70)
print("SUMMARY")
print("="*70)

all_found = True
for key, found in improvements.items():
    status = "✅" if found else "❌"
    print(f"{status} {key}")
    if not found:
        all_found = False

print("\n" + "="*70)
if all_found:
    print("SUCCESS: All improvements verified in the document!")
else:
    print("WARNING: Some improvements may be missing or named differently")
print("="*70)

print(f"\nTotal paragraphs in document: {len(doc.paragraphs)}")
