"""Part 2: Spelling, merged words, missing spaces"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
import re

print("Loading KDP-final.docx...")
doc = Document(r'public\reference\book\KDP-final.docx')
print(f"Loaded: {len(doc.paragraphs)} paragraphs")

out = open('audit_p2_results.txt', 'w', encoding='utf-8')

def w(text=""):
    out.write(text + "\n")

w("PART 2: SPELLING & WORD MERGING - KDP-final.docx")
w("=" * 60)

# Merged word detection: lowercase immediately followed by uppercase mid-word
legit_camel = {
    'MHz','GHz','mATX','eReader','iCloud','macOS','iOS','iPadOS',
    'watchOS','tvOS','iPhone','iPad','iPod','BitLocker','PowerShell',
    'VirtualBox','TeamViewer','AnyDesk','JavaScript','GitHub',
    'YouTube','WiFi','FireOS','ChromeOS','exFAT','CardDAV','CalDAV',
    'ActiveSync','FileVault','AirDrop','TimeMachine','OneNote',
    'OneDrive','InPrivate','QuickBooks','AutoRun','AutoPlay',
    'McAfee','WinRE','ReadyBoost','HyperV','EasyBCD','PageFile',
    'DevOps','FortiGate','NetBIOS','IntelliSense','TypeScript',
    'AppLocker','CrowdStrike','SentinelOne','NordVPN','OpenVPN',
    'WireGuard','AnyConnect','FileZilla','PuTTY','WinSCP',
    'TechNet','SuperSpeed','DisplayPort','MicroSD','CompTIA',
    'ServiceNow','ConnectWise','SolarWinds','CloudFlare',
}

merged_found = {}
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if not text.strip():
        continue
    # Find camelCase-like patterns in running text
    for match in re.finditer(r'\b([a-z]+[A-Z][a-z]+\w*)\b', text):
        word = match.group()
        if word not in legit_camel and word not in merged_found:
            merged_found[word] = (i, text[:70])
    # Also find uppercase followed by lowercase without space: "WordWord"
    for match in re.finditer(r'([A-Z][a-z]+)([A-Z][a-z]+)', text):
        combo = match.group()
        if combo not in legit_camel and combo not in merged_found:
            # Skip if it's a known proper noun
            if len(combo) > 8:  # Only flag longer merges
                merged_found[combo] = (i, text[:70])

w(f"\n[Merged Words] Found: {len(merged_found)}")
for word, (idx, ctx) in list(merged_found.items())[:30]:
    w(f"  '{word}' (para {idx}): {ctx}")

# Missing space after period (sentence.Next without space)
missing_space = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # lowercase letter, period, uppercase letter with no space
    for m in re.finditer(r'([a-z])\.([A-Z])', text):
        context = text[max(0,m.start()-10):m.end()+20]
        missing_space.append((i, m.group(), context))

w(f"\n[Missing Space After Period] Found: {len(missing_space)}")
for idx, match, ctx in missing_space[:20]:
    w(f"  Para {idx}: '{match}' in ...{ctx}...")

# Common misspellings check
misspell_patterns = [
    (r'\brecieve\b', 'receive'), (r'\boccured\b', 'occurred'),
    (r'\bseperete\b', 'separate'), (r'\bneccessary\b', 'necessary'),
    (r'\benviroment\b', 'environment'), (r'\bperformace\b', 'performance'),
    (r'\bconfiguartion\b', 'configuration'), (r'\bauthentiction\b', 'authentication'),
    (r'\btroubleshooing\b', 'troubleshooting'), (r'\bencryptoin\b', 'encryption'),
    (r'\bconnecton\b', 'connection'), (r'\bnetowrk\b', 'network'),
    (r'\boperting\b', 'operating'), (r'\bsytems\b', 'systems'),
    (r'\bmanagment\b', 'management'), (r'\bvirtualiztion\b', 'virtualization'),
    (r'\btechnolgy\b', 'technology'), (r'\binfromation\b', 'information'),
    (r'\bsercurity\b', 'security'), (r'\bcompatabile\b', 'compatible'),
    (r'\bthier\b', 'their'), (r'\bteh\b', 'the'),
    (r'\bwich\b', 'which'), (r'\bbeacuse\b', 'because'),
    (r'\bwindwos\b', 'windows'), (r'\bmotherbaord\b', 'motherboard'),
    (r'\bprocesssor\b', 'processor'), (r'\bconnectivty\b', 'connectivity'),
    (r'\bencyrption\b', 'encryption'), (r'\bauthetication\b', 'authentication'),
    (r'\bpermisions\b', 'permissions'), (r'\btroubleshoting\b', 'troubleshooting'),
    (r'\bvulnerabilty\b', 'vulnerability'), (r'\bfirewal\b', 'firewall'),
    (r'\bencryptio\b', 'encryption'), (r'\bprotocl\b', 'protocol'),
]

spelling_issues = []
for para in doc.paragraphs:
    text = para.text
    for pattern, correct in misspell_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            spelling_issues.append((re.search(pattern, text, re.IGNORECASE).group(), correct, text[:60]))

w(f"\n[Misspellings] Found: {len(spelling_issues)}")
for wrong, correct, ctx in spelling_issues[:15]:
    w(f"  '{wrong}' -> '{correct}': {ctx}")

# Repeated words (the the, is is, etc.)
repeated = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    for m in re.finditer(r'\b(\w{3,})\s+\1\b', text, re.IGNORECASE):
        repeated.append((i, m.group(), text[:60]))

w(f"\n[Repeated Words] Found: {len(repeated)}")
for idx, match, ctx in repeated[:15]:
    w(f"  Para {idx}: '{match}' in: {ctx}")

out.close()
print("Done - audit_p2_results.txt")
